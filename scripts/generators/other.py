"""Miscellaneous form generators for KFSYSCC IRB.

Generates SF076 form from config dict.
"""
import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# SF076 — 閱卷複印申請登記表 (v1)
# ---------------------------------------------------------------------------


def generate_sf076(config, output_dir):
    """Generate SF076: document review and copy request form."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "閱卷複印申請登記表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Simple header: IRB no and title only
    tbl_h = doc.add_table(rows=2, cols=2)
    tbl_h.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl_h.rows[0].cells[0], "KFSYSCC-IRB編號", bold=True, size=11)
    add_ct(tbl_h.rows[0].cells[1], config["study"]["irb_no"], size=11)
    add_ct(tbl_h.rows[1].cells[0], "計畫名稱", bold=True, size=11)
    add_ct(tbl_h.rows[1].cells[1], config["study"]["title_zh"], size=11)
    apply_tb(tbl_h)
    doc.add_paragraph()

    # Applicant info table
    tbl_a = doc.add_table(rows=4, cols=2)
    tbl_a.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl_a.rows[0].cells[0], "申請人姓名", bold=True, size=10)
    add_ct(tbl_a.rows[0].cells[1], pi["name"], size=10)
    add_ct(tbl_a.rows[1].cells[0], "單位", bold=True, size=10)
    add_ct(tbl_a.rows[1].cells[1], pi.get("dept", ""), size=10)
    add_ct(tbl_a.rows[2].cells[0], "電話", bold=True, size=10)
    add_ct(tbl_a.rows[2].cells[1], pi.get("phone", ""), size=10)
    add_ct(tbl_a.rows[3].cells[0], "申請日期", bold=True, size=10)
    add_ct(tbl_a.rows[3].cells[1], "", size=10)
    apply_tb(tbl_a)
    doc.add_paragraph()

    # Requested documents table
    doc_cols = ["項次", "文件名稱", "份數", "備註"]
    tbl_d = doc.add_table(rows=1 + 5, cols=4)
    tbl_d.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(doc_cols):
        cell = tbl_d.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # 5 empty rows
    for ri in range(1, 6):
        add_ct(tbl_d.rows[ri].cells[0], str(ri), size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl_d.rows[ri].cells[1], "", size=10)
        add_ct(tbl_d.rows[ri].cells[2], "", size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl_d.rows[ri].cells[3], "", size=10)

    apply_tb(tbl_d)
    doc.add_paragraph()

    # Purpose
    add_p(doc, "申請目的：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿",
          size=11, sa=Pt(8))

    # Applicant signature
    add_p(doc, "申請人簽名：＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=11, sa=Pt(8))

    # IRB staff approval
    add_p(doc, "IRB承辦人：＿＿＿＿＿＿＿＿　核准日期：＿＿＿＿年＿＿月＿＿日",
          size=11, sa=Pt(12))

    # Footer
    add_footer(doc, "1", "IRB.SF076", "2022/09/01")

    path = os.path.join(output_dir, "SF076_閱卷複印申請登記表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF076": generate_sf076,
}
