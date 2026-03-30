"""Re-review (複審) form generators for KFSYSCC IRB.

Generates SF019 from config dict.
"""
import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# SF019 — 複審案申請表 v4, 2018/01/02
# ---------------------------------------------------------------------------


def generate_sf019(config, output_dir):
    """Generate SF019: re-review application form."""
    doc = init_doc(sz=12)

    # Title block
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(0), sb=Pt(12))
    add_p(doc, "複審案申請表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(0), sb=Pt(6))

    # Header block
    add_header(doc, config, inc_proj=True)

    # ------------------------------------------------------------------
    # Section 一、原審查案資料
    # ------------------------------------------------------------------
    add_p(doc, "一、原審查案資料：", bold=True, size=12, sa=Pt(8), sb=Pt(4))

    tbl1 = doc.add_table(rows=4, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_ct(tbl1.rows[0].cells[0], "IRB編號", bold=True, size=11)
    add_ct(tbl1.rows[0].cells[1], config["study"]["irb_no"], size=11)

    add_ct(tbl1.rows[1].cells[0], "計畫名稱", bold=True, size=11)
    cell_title = tbl1.rows[1].cells[1]
    p = cell_title.paragraphs[0]
    p.clear()
    run = p.add_run(f"（中文）{config['study']['title_zh']}")
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell_title.add_paragraph()
    run2 = p2.add_run(f"（英文）{config['study']['title_en']}")
    set_run_font(run2, size=10)
    p2.paragraph_format.space_after = Pt(2)

    add_ct(tbl1.rows[2].cells[0], "主持人", bold=True, size=11)
    add_ct(tbl1.rows[2].cells[1], config["pi"]["name"], size=11)

    add_ct(tbl1.rows[3].cells[0], "單位／職稱", bold=True, size=11)
    add_ct(tbl1.rows[3].cells[1], config["pi"]["dept"], size=11)

    apply_tb(tbl1)

    # 原審查類別
    add_p(doc, "", sa=Pt(2), sb=Pt(0))
    add_p(doc, "原審查類別：  □ 新案　□ 修正案　□ 期中審查　□ 其他＿＿＿＿",
          size=12, sa=Pt(4), sb=Pt(2))
    add_p(doc, "原審查日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(4))

    # ------------------------------------------------------------------
    # Section 二、審查意見回覆
    # ------------------------------------------------------------------
    add_p(doc, "二、審查意見回覆：", bold=True, size=12, sa=Pt(8), sb=Pt(4))
    add_p(doc, "請逐項回覆委員會之審查意見", size=12, sa=Pt(2), sb=Pt(4))

    tbl2 = doc.add_table(rows=6, cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["項次", "審查意見", "回覆說明"]
    for ci, h in enumerate(headers):
        cell = tbl2.rows[0].cells[ci]
        add_ct(cell, h, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # 5 numbered empty rows
    for ri in range(1, 6):
        add_ct(tbl2.rows[ri].cells[0], str(ri), size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl2.rows[ri].cells[1], "", size=10)
        add_ct(tbl2.rows[ri].cells[2], "", size=10)

    apply_tb(tbl2)

    # ------------------------------------------------------------------
    # Section 三、修正內容
    # ------------------------------------------------------------------
    add_p(doc, "三、修正內容：", bold=True, size=12, sa=Pt(8), sb=Pt(4))
    add_p(doc, "請以劃線方式標示修正處，並附上修正後之相關文件",
          size=12, sa=Pt(2), sb=Pt(4))

    items = [
        "□ 修正後研究計畫書",
        "□ 修正後受試者同意書",
        "□ 修正後個案報告表",
        "□ 其他（請說明）：＿＿＿＿＿＿＿＿＿＿",
    ]
    for item in items:
        add_p(doc, f"  {item}", size=12, sa=Pt(2), sb=Pt(2))

    # ------------------------------------------------------------------
    # Section 四、主持人聲明
    # ------------------------------------------------------------------
    add_p(doc, "四、主持人聲明：", bold=True, size=12, sa=Pt(8), sb=Pt(4))
    add_p(doc, "本人聲明以上回覆內容均屬實，並已依審查意見修正相關文件。",
          size=12, sa=Pt(4), sb=Pt(6))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(6))

    # Note at bottom
    add_p(doc, "")
    add_p(doc, "一式二份（正本一份、影本一份），連同申請表及修正文件送交IRB",
          size=10, sa=Pt(6), sb=Pt(2))

    # Footer
    add_footer(doc, "4", "IRB.SF019", "2018/01/02")

    path = os.path.join(output_dir, "SF019_複審案申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# ALL_GENERATORS registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF019": generate_sf019,
}
