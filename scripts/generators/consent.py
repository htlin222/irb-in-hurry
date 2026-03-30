"""Consent / waiver / review-type form generators for KFSYSCC IRB.

Generates: SF003, SF004, SF005, SF062, SF063, SF075, SF090, SF091, SF092.
"""
import os

from scripts.docx_utils import (
    init_doc, set_cell_shading, add_p, add_ct, apply_tb,
    add_header, add_footer, check, set_run_font,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _title_block(doc, subtitle):
    """Add standard two-line title block used by consent/review forms."""
    add_p(doc, "和信治癌中心醫院 人體試驗委員會",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(0), sb=Pt(12))
    add_p(doc, subtitle,
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(0), sb=Pt(6))


def _consent_title_block(doc, subtitle):
    """Add consent-form style title (hospital name + form subtitle)."""
    add_p(doc, "和信治癌中心醫院",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(0), sb=Pt(12))
    add_p(doc, subtitle,
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(0), sb=Pt(6))


def _study_info_block(doc, config):
    """Add IRB number and study title info for consent forms."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_ct(tbl.rows[0].cells[0], "IRB編號", bold=True, size=11)
    add_ct(tbl.rows[0].cells[1], config["study"]["irb_no"], size=11)

    add_ct(tbl.rows[1].cells[0], "計畫名稱", bold=True, size=11)
    cell = tbl.rows[1].cells[1]
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f"（中文）{config['study']['title_zh']}")
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(f"（英文）{config['study']['title_en']}")
    set_run_font(run2, size=10)
    p2.paragraph_format.space_after = Pt(2)

    apply_tb(tbl)
    doc.add_paragraph()


def _consent_contact_section(doc, config, section_num):
    """Add standard contact section for consent forms."""
    pi_name = config["pi"]["name"]
    pi_phone = config["pi"].get("phone", "＿＿＿＿＿＿＿＿")
    pi_email = config["pi"].get("email", "＿＿＿＿＿＿＿＿")

    add_p(doc, f"{section_num}、聯絡人", bold=True, size=12, sa=Pt(8), sb=Pt(2))
    add_p(doc, f"如果您對本研究有任何疑問，請聯繫：", size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, f"  計畫主持人：{pi_name}", size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, f"  電話：{pi_phone}", size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, f"  電子郵件：{pi_email}", size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "  人體試驗委員會聯絡電話：(02) 2897-0011 分機3952",
          size=12, sa=Pt(2), sb=Pt(2))


def _consent_rights_section(doc, section_num):
    """Add standard participant rights section."""
    add_p(doc, f"{section_num}、受試者權益", bold=True, size=12, sa=Pt(8), sb=Pt(2))
    add_p(doc, "如您對研究參與之權益有疑問，可聯繫本院人體試驗委員會。",
          size=12, sa=Pt(2), sb=Pt(2))


def _consent_signature_block(doc):
    """Add full consent form signature block."""
    add_p(doc, "")
    add_p(doc, "受試者同意聲明", bold=True, size=12, sa=Pt(8), sb=Pt(4))
    add_p(doc, "本人已詳閱以上說明，並有機會詢問相關問題，本人同意參加本研究。",
          size=12, sa=Pt(2), sb=Pt(6))
    add_p(doc, "受試者簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "法定代理人簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日　（如適用）",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "關係：＿＿＿＿＿＿",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "見證人簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "計畫主持人/研究人員簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))


def _signature_block(doc):
    """Add PI signature + date line."""
    add_p(doc, "")
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(0), sb=Pt(6))


def _save(doc, output_dir, filename):
    """Save document and return path."""
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF003 - 簡易審查範圍檢核表  v4, 2018/01/02
# ---------------------------------------------------------------------------

def generate_sf003(config, output_dir):
    """Generate SF003 expedited review scope checklist."""
    doc = init_doc()
    _title_block(doc, "簡易審查範圍檢核表")
    add_header(doc, config, inc_proj=True)

    add_p(doc, "本計畫符合下列簡易審查之範圍（請勾選）：", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    is_retro = config["study"]["type"] == "retrospective"

    items = [
        (False, "1. 已核准之藥品或醫療器材，於新適應症或新用途之臨床試驗"),
        (False, "2. 利用無侵入性方式收集之生物檢體供研究用途"),
        (False, "3. 以非侵入性方式收集之臨床資料供研究用途"),
        (is_retro, "4. 既有資料、文件、紀錄或病理檢體之研究"),
        (False, "5. 聯邦政府或機構為研究、公共利益或公共服務計畫所收集之資料"),
        (False, "6. 食品之品質評估及消費者接受度之研究"),
        (False, "7. 其他經主管機關公告之研究"),
    ]

    for condition, text in items:
        add_p(doc, f"  {check(condition)} {text}", size=12, sa=Pt(2), sb=Pt(2))

    add_p(doc, "")
    add_p(doc, "備註說明：", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", size=12, sa=Pt(2), sb=Pt(2))

    _signature_block(doc)
    add_footer(doc, "4", "IRB.SF003", "2018/01/02")

    return _save(doc, output_dir, "SF003_簡易審查範圍檢核表.docx")


# ---------------------------------------------------------------------------
# SF004 - 免予審查範圍檢核表  v2, 2016/11/07
# ---------------------------------------------------------------------------

def generate_sf004(config, output_dir):
    """Generate SF004 exempt review scope checklist."""
    doc = init_doc()
    _title_block(doc, "免予審查範圍檢核表")
    add_header(doc, config, inc_proj=True)

    add_p(doc, "本計畫符合下列免予審查之範圍（請勾選）：", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    is_retro = config["study"]["type"] == "retrospective"

    items = [
        (False, "1. 於一般教學環境中進行之教育策略研究"),
        (False, "2. 教育測驗、訪談程序、問卷調查或公開行為觀察之研究"),
        (is_retro, "3. 利用既有之資料、文件、紀錄、病理檢體或診斷檢體之研究，且資料來源為公開或去識別化"),
        (False, "4. 公職人員或公職候選人之研究"),
        (False, "5. 食品之品質評估研究"),
        (False, "6. 其他"),
    ]

    for condition, text in items:
        add_p(doc, f"  {check(condition)} {text}", size=12, sa=Pt(2), sb=Pt(2))

    _signature_block(doc)
    add_footer(doc, "2", "IRB.SF004", "2016/11/07")

    return _save(doc, output_dir, "SF004_免予審查範圍檢核表.docx")


# ---------------------------------------------------------------------------
# SF005 - 免取得知情同意檢核表  v2, 2016/11/07
# ---------------------------------------------------------------------------

def generate_sf005(config, output_dir):
    """Generate SF005 waiver of consent checklist."""
    doc = init_doc()
    _title_block(doc, "免取得知情同意檢核表")
    add_header(doc, config, inc_proj=True)

    add_p(doc, "申請免取得受試者知情同意書之理由（請勾選符合之項目）：",
          bold=True, size=12, sa=Pt(6), sb=Pt(6))

    # All four criteria must be met for waiver to be valid
    items = [
        "1. 本研究對受試者所產生之風險不超過最低風險",
        "2. 免取得知情同意對受試者之權益不致產生不良影響",
        "3. 因研究設計之需要，無法事前取得受試者之知情同意",
        "4. 於適當情況下，受試者於研究參與後將會被提供額外之相關資訊",
    ]

    for text in items:
        add_p(doc, f"  {check(True)} {text}", size=12, sa=Pt(2), sb=Pt(2))

    # Explanation section
    data_period = config.get("dates", {}).get("data_period", "＿＿年至＿＿年間")
    explanation = (
        f"本研究為回溯性病歷審查研究，僅回顧分析{data_period}之"
        "既有去識別化病歷資料，不涉及任何介入性措施或直接接觸受試者。"
        "研究資料均已去識別化處理，無法連結至個別受試者。"
        "基於研究性質，事前取得每位受試者之知情同意實務上不可行且不必要。"
    )

    add_p(doc, "")
    add_p(doc, "說明：", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, explanation, size=12, sa=Pt(2), sb=Pt(2))

    _signature_block(doc)
    add_footer(doc, "2", "IRB.SF005", "2016/11/07")

    return _save(doc, output_dir, "SF005_免取得知情同意檢核表.docx")


# ---------------------------------------------------------------------------
# SF062 — 受試者同意書（臨床研究）v4, 2023/11/01
# ---------------------------------------------------------------------------

def generate_sf062(config, output_dir):
    """Generate SF062 consent form for non-drug clinical research."""
    doc = init_doc(sz=12)
    _consent_title_block(doc, "受試者同意書（臨床研究）")
    _study_info_block(doc, config)

    sections = [
        ("一、研究說明",
         "您被邀請參加一項臨床研究。在您決定是否參加之前，請仔細閱讀以下內容。"
         "研究人員會向您說明本研究的相關資訊，您可以隨時提出問題。"),
        ("二、研究目的",
         "（請說明本研究之目的及預期成果）"),
        ("三、研究方法與程序",
         "（請說明研究之方法、流程、持續時間、受試者須配合之事項）"),
        ("四、可能的風險與不適",
         "（請說明受試者可能面臨之風險、副作用或不適，以及因應措施）"),
        ("五、可能的益處",
         "（請說明受試者及社會可能獲得之益處）"),
        ("六、替代方案",
         "（請說明不參與本研究時可選擇之其他治療或處置方式）"),
        ("七、費用與補償",
         "本研究不額外收取費用。（請說明受試者參與研究之交通費補助或其他補償方式）"),
        ("八、保密性",
         "您的個人資料將依法受到保護。研究人員將採取適當之保密措施，"
         "確保您的個人隱私不會被洩漏。研究結果發表時，將不會揭露您的身分。"
         "但依法規要求，主管機關、人體試驗委員會得於必要時查閱您的研究資料。"),
        ("九、自願參加與退出",
         "您的參與完全自願。您可以在任何時候退出研究，不會影響您的醫療照護品質。"
         "您的退出不需要任何理由，也不會影響您與醫療人員之間的關係。"),
    ]

    for title, content in sections:
        add_p(doc, title, bold=True, size=12, sa=Pt(8), sb=Pt(2))
        add_p(doc, content, size=12, sa=Pt(2), sb=Pt(2))

    _consent_contact_section(doc, config, "十")
    _consent_rights_section(doc, "十一")
    _consent_signature_block(doc)

    add_footer(doc, "4", "IRB.SF062", "2023/11/01")
    return _save(doc, output_dir, "SF062_受試者同意書（臨床研究）.docx")


# ---------------------------------------------------------------------------
# SF063 — 受試者同意書（臨床試驗）v11, 2023/11/01
# ---------------------------------------------------------------------------

def generate_sf063(config, output_dir):
    """Generate SF063 consent form for drug/device clinical trial."""
    doc = init_doc(sz=12)
    _consent_title_block(doc, "受試者同意書（臨床試驗）")
    _study_info_block(doc, config)

    sections = [
        ("一、研究說明",
         "您被邀請參加一項臨床試驗。在您決定是否參加之前，請仔細閱讀以下內容。"
         "研究人員會向您說明本試驗的相關資訊，您可以隨時提出問題。"),
        ("二、研究目的",
         "（請說明本臨床試驗之目的及預期成果）"),
        ("三、試驗藥品/醫療器材名稱",
         "（請填寫試驗藥品或醫療器材之名稱、劑量、給藥方式）"),
        ("四、研究方法與程序",
         "（請說明試驗之方法、流程、持續時間、受試者須配合之事項，包括隨機分派及雙盲設計等）"),
        ("五、可能的風險與不適",
         "（請說明受試者可能面臨之風險、副作用或不適，以及因應措施）"),
        ("六、可能的益處",
         "（請說明受試者及社會可能獲得之益處）"),
        ("七、替代方案",
         "（請說明不參與本試驗時可選擇之其他治療或處置方式）"),
        ("八、藥品管理與退還",
         "試驗期間所提供之試驗用藥品，應依研究人員指示使用，不得轉讓他人。"
         "試驗結束或退出時，受試者應將未使用之藥品及空瓶退還研究人員。"),
        ("九、懷孕與避孕注意事項",
         "（請說明試驗期間之避孕要求、懷孕通報程序及對胎兒之潛在風險）"),
        ("十、費用與補償",
         "本試驗不額外收取費用。（請說明受試者參與試驗之交通費補助或其他補償方式）"),
        ("十一、試驗相關傷害之補償",
         "若您因參與本試驗而受到傷害，將依相關法規提供適當之醫療照護與補償。"
         "（請說明試驗相關傷害之醫療處理及補償機制）"),
        ("十二、保密性",
         "您的個人資料將依法受到保護。研究人員將採取適當之保密措施，"
         "確保您的個人隱私不會被洩漏。研究結果發表時，將不會揭露您的身分。"
         "但依法規要求，主管機關、人體試驗委員會及試驗委託者得於必要時查閱您的研究資料。"),
        ("十三、自願參加與退出",
         "您的參與完全自願。您可以在任何時候退出試驗，不會影響您的醫療照護品質。"
         "您的退出不需要任何理由，也不會影響您與醫療人員之間的關係。"),
    ]

    for title, content in sections:
        add_p(doc, title, bold=True, size=12, sa=Pt(8), sb=Pt(2))
        add_p(doc, content, size=12, sa=Pt(2), sb=Pt(2))

    _consent_contact_section(doc, config, "十四")
    _consent_rights_section(doc, "十五")
    _consent_signature_block(doc)

    add_footer(doc, "11", "IRB.SF063", "2023/11/01")
    return _save(doc, output_dir, "SF063_受試者同意書（臨床試驗）.docx")


# ---------------------------------------------------------------------------
# SF075 — 受試者同意書（基因研究）v3, 2023/11/01
# ---------------------------------------------------------------------------

def generate_sf075(config, output_dir):
    """Generate SF075 consent form for genetic research."""
    doc = init_doc(sz=12)
    _consent_title_block(doc, "受試者同意書（基因研究）")
    _study_info_block(doc, config)

    sections = [
        ("一、研究說明",
         "您被邀請參加一項基因研究。在您決定是否參加之前，請仔細閱讀以下內容。"
         "研究人員會向您說明本研究的相關資訊，您可以隨時提出問題。"),
        ("二、研究目的",
         "（請說明本基因研究之目的及預期成果）"),
        ("三、研究方法與程序",
         "（請說明研究之方法、流程、持續時間、受試者須配合之事項）"),
        ("四、基因檢測說明",
         "（請說明本研究所涉及之基因檢測項目、檢測方法，以及基因檢測結果之意義與限制）"),
        ("五、檢體之保存與使用",
         "您的檢體將用於本研究之基因分析。\n"
         "  保存期限：（請說明檢體保存之期限）\n"
         "  使用範圍：（請說明檢體使用之範圍，是否用於未來相關研究）\n"
         "  銷毀方式：（請說明研究結束後檢體之銷毀方式）"),
        ("六、可能的風險與不適",
         "（請說明受試者可能面臨之風險、副作用或不適，以及因應措施）"),
        ("七、可能的益處",
         "（請說明受試者及社會可能獲得之益處）"),
        ("八、替代方案",
         "（請說明不參與本研究時可選擇之其他方案）"),
        ("九、費用與補償",
         "本研究不額外收取費用。（請說明受試者參與研究之交通費補助或其他補償方式）"),
        ("十、基因資訊之保密",
         "您的基因資訊屬於高度敏感之個人資料，研究人員將採取嚴格之保密措施加以保護。"
         "基因資料將以編碼方式儲存，與您的個人身分資料分開保管。"
         "未經您的同意，研究人員不會將您的基因資訊提供給第三方，"
         "包括保險公司、雇主或其他可能影響您權益之機構。"
         "研究結果發表時，將不會揭露您的身分或可供辨識之基因資訊。"),
        ("十一、自願參加與退出",
         "您的參與完全自願。您可以在任何時候退出研究，不會影響您的醫療照護品質。"
         "若您決定退出，您可以選擇要求銷毀已採集之檢體及基因資料。"),
        ("十二、研究結果之告知",
         "  □ 願意被告知研究結果\n"
         "  □ 不願意被告知研究結果"),
    ]

    for title, content in sections:
        add_p(doc, title, bold=True, size=12, sa=Pt(8), sb=Pt(2))
        add_p(doc, content, size=12, sa=Pt(2), sb=Pt(2))

    _consent_contact_section(doc, config, "十三")
    _consent_rights_section(doc, "十四")
    _consent_signature_block(doc)

    add_footer(doc, "3", "IRB.SF075", "2023/11/01")
    return _save(doc, output_dir, "SF075_受試者同意書（基因研究）.docx")


# ---------------------------------------------------------------------------
# SF090 — 受試者同意書（藥品試驗）v3, 2025/03/03
# ---------------------------------------------------------------------------

def generate_sf090(config, output_dir):
    """Generate SF090 consent form for drug trial (TFDA enhanced)."""
    doc = init_doc(sz=12)
    _consent_title_block(doc, "受試者同意書（藥品試驗）")
    _study_info_block(doc, config)

    add_p(doc, "依藥品優良臨床試驗準則（ICH-GCP）規範",
          bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
          sa=Pt(0), sb=Pt(6))

    sections = [
        ("一、研究說明",
         "您被邀請參加一項藥品臨床試驗。在您決定是否參加之前，請仔細閱讀以下內容。"
         "研究人員會向您說明本試驗的相關資訊，您可以隨時提出問題。"),
        ("二、研究目的",
         "（請說明本藥品臨床試驗之目的及預期成果）"),
        ("三、試驗藥品名稱",
         "（請填寫試驗藥品之名稱、劑量、給藥方式、給藥頻率）"),
        ("四、研究方法與程序",
         "（請說明試驗之方法、流程、持續時間、受試者須配合之事項，"
         "包括隨機分派、雙盲設計及安慰劑使用等）"),
        ("五、藥品動力學檢查",
         "（如適用，請說明藥品動力學檢查之目的、方法、"
         "需採集之血液或尿液樣本數量及時間點）"),
        ("六、可能的風險與不適",
         "（請說明受試者可能面臨之風險、副作用或不適，以及因應措施）"),
        ("七、可能的益處",
         "（請說明受試者及社會可能獲得之益處）"),
        ("八、替代方案",
         "（請說明不參與本試驗時可選擇之其他治療方式）"),
        ("九、藥品管理與退還",
         "試驗期間所提供之試驗用藥品，應依研究人員指示使用，不得轉讓他人。"
         "試驗結束或退出時，受試者應將未使用之藥品及空瓶退還研究人員。"),
        ("十、懷孕與避孕注意事項",
         "（請說明試驗期間之避孕要求、懷孕通報程序及對胎兒之潛在風險）"),
        ("十一、費用與補償",
         "本試驗不額外收取費用。（請說明受試者參與試驗之交通費補助或其他補償方式）"),
        ("十二、試驗相關傷害之補償",
         "若您因參與本試驗而受到傷害，將依相關法規提供適當之醫療照護與補償。"
         "（請說明試驗相關傷害之醫療處理及補償機制）"),
        ("十三、試驗用藥品之保險",
         "（請說明試驗委託者是否為受試者投保臨床試驗責任保險，以及保險之範圍與金額）"),
        ("十四、保密性",
         "您的個人資料將依法受到保護。研究人員將採取適當之保密措施，"
         "確保您的個人隱私不會被洩漏。研究結果發表時，將不會揭露您的身分。"
         "但依法規要求，主管機關、人體試驗委員會及試驗委託者得於必要時查閱您的研究資料。"),
        ("十五、自願參加與退出",
         "您的參與完全自願。您可以在任何時候退出試驗，不會影響您的醫療照護品質。"
         "您的退出不需要任何理由，也不會影響您與醫療人員之間的關係。"),
    ]

    for title, content in sections:
        add_p(doc, title, bold=True, size=12, sa=Pt(8), sb=Pt(2))
        add_p(doc, content, size=12, sa=Pt(2), sb=Pt(2))

    _consent_contact_section(doc, config, "十六")
    _consent_rights_section(doc, "十七")
    _consent_signature_block(doc)

    add_footer(doc, "3", "IRB.SF090", "2025/03/03")
    return _save(doc, output_dir, "SF090_受試者同意書（藥品試驗）.docx")


# ---------------------------------------------------------------------------
# SF091 — 受試者同意書（兒童）v2, 2023/11/01
# ---------------------------------------------------------------------------

def generate_sf091(config, output_dir):
    """Generate SF091 consent/assent form for pediatric participants."""
    doc = init_doc(sz=12)
    _consent_title_block(doc, "受試者同意書（兒童/青少年）")
    _study_info_block(doc, config)

    # ---- Part A: Parent/Guardian consent ----
    add_p(doc, "【第一部分：法定代理人同意書】",
          bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER,
          sa=Pt(8), sb=Pt(6))

    sections = [
        ("一、研究說明",
         "您的孩子被邀請參加一項研究。在您決定是否讓孩子參加之前，"
         "請仔細閱讀以下內容。研究人員會向您說明本研究的相關資訊，您可以隨時提出問題。"),
        ("二、研究目的",
         "（請說明本研究之目的及預期成果）"),
        ("三、研究方法與程序",
         "（請說明研究之方法、流程、持續時間、孩子須配合之事項）"),
        ("四、可能的風險與不適",
         "（請說明孩子可能面臨之風險、副作用或不適，以及因應措施）"),
        ("五、可能的益處",
         "（請說明孩子及社會可能獲得之益處）"),
        ("六、替代方案",
         "（請說明不參與本研究時可選擇之其他方案）"),
        ("七、費用與補償",
         "本研究不額外收取費用。（請說明參與研究之交通費補助或其他補償方式）"),
        ("八、保密性",
         "您孩子的個人資料將依法受到保護。研究人員將採取適當之保密措施，"
         "確保孩子的個人隱私不會被洩漏。研究結果發表時，將不會揭露孩子的身分。"),
        ("九、自願參加與退出",
         "您孩子的參與完全自願。您可以在任何時候讓孩子退出研究，不會影響孩子的醫療照護品質。"),
    ]

    for title, content in sections:
        add_p(doc, title, bold=True, size=12, sa=Pt(8), sb=Pt(2))
        add_p(doc, content, size=12, sa=Pt(2), sb=Pt(2))

    _consent_contact_section(doc, config, "十")
    _consent_rights_section(doc, "十一")

    # Parent/Guardian signature
    add_p(doc, "")
    add_p(doc, "法定代理人同意聲明", bold=True, size=12, sa=Pt(8), sb=Pt(4))
    add_p(doc, "本人已詳閱以上說明，並有機會詢問相關問題，本人同意讓孩子參加本研究。",
          size=12, sa=Pt(2), sb=Pt(6))
    add_p(doc, "法定代理人（父/母/監護人）簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "與受試者之關係：＿＿＿＿＿＿",
          size=12, sa=Pt(2), sb=Pt(2))

    # ---- Part B: Child assent ----
    add_p(doc, "")
    add_p(doc, "【第二部分：兒童/青少年同意書（簡易版）】",
          bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER,
          sa=Pt(12), sb=Pt(6))

    add_p(doc, "親愛的小朋友/同學：", bold=True, size=12, sa=Pt(6), sb=Pt(4))
    add_p(doc, "這份文件是要告訴你關於一個研究。"
          "我們想請你幫忙參加這個研究，但是你可以自己決定要不要參加。",
          size=12, sa=Pt(2), sb=Pt(4))
    add_p(doc, "這個研究是什麼？", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, "（請以兒童/青少年可理解之語言說明研究內容）",
          size=12, sa=Pt(2), sb=Pt(4))
    add_p(doc, "參加研究會怎樣？", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, "（請以簡單語言說明研究過程中孩子需要做的事情）",
          size=12, sa=Pt(2), sb=Pt(4))
    add_p(doc, "會不會不舒服？", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, "（請以簡單語言說明可能的不適）",
          size=12, sa=Pt(2), sb=Pt(4))
    add_p(doc, "有什麼好處？", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, "（請以簡單語言說明可能的好處）",
          size=12, sa=Pt(2), sb=Pt(4))
    add_p(doc, "你可以自己決定！", bold=True, size=12, sa=Pt(6), sb=Pt(2))
    add_p(doc, "如果你不想參加，沒有關係，不會有人生氣。"
          "就算你已經說要參加了，之後不想繼續也可以隨時退出。",
          size=12, sa=Pt(2), sb=Pt(4))

    # Child assent signature
    add_p(doc, "")
    add_p(doc, "受試者（兒童/青少年）同意聲明", bold=True, size=12, sa=Pt(8), sb=Pt(4))
    add_p(doc, "有人已經告訴我這個研究是什麼，我也可以問問題。我願意參加這個研究。",
          size=12, sa=Pt(2), sb=Pt(6))
    add_p(doc, "受試者（兒童/青少年）簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "見證人簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))
    add_p(doc, "計畫主持人/研究人員簽名：＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=12, sa=Pt(2), sb=Pt(2))

    add_footer(doc, "2", "IRB.SF091", "2023/11/01")
    return _save(doc, output_dir, "SF091_受試者同意書（兒童）.docx")


# ---------------------------------------------------------------------------
# SF092 — 受試者同意書（醫療器材）v2, 2023/11/01
# ---------------------------------------------------------------------------

def generate_sf092(config, output_dir):
    """Generate SF092 consent form for medical device trial."""
    doc = init_doc(sz=12)
    _consent_title_block(doc, "受試者同意書（醫療器材）")
    _study_info_block(doc, config)

    sections = [
        ("一、研究說明",
         "您被邀請參加一項醫療器材臨床試驗。在您決定是否參加之前，請仔細閱讀以下內容。"
         "研究人員會向您說明本試驗的相關資訊，您可以隨時提出問題。"),
        ("二、研究目的",
         "（請說明本醫療器材試驗之目的及預期成果）"),
        ("三、醫療器材名稱與說明",
         "（請填寫醫療器材之名稱、型號、許可證字號或專案核准文號，"
         "以及器材之功能與用途說明）"),
        ("四、器材植入/使用方式",
         "（請說明醫療器材之植入或使用方式、使用頻率、使用期間，"
         "以及受試者須配合之事項）"),
        ("五、研究方法與程序",
         "（請說明試驗之方法、流程、持續時間、受試者須配合之事項）"),
        ("六、可能的風險與不適",
         "（請說明受試者可能面臨之風險、副作用或不適，以及因應措施）"),
        ("七、器材故障或失效之處理",
         "（請說明醫療器材發生故障或失效時之處理方式、通報程序，"
         "以及對受試者之影響與因應措施）"),
        ("八、可能的益處",
         "（請說明受試者及社會可能獲得之益處）"),
        ("九、替代方案",
         "（請說明不參與本試驗時可選擇之其他治療或處置方式）"),
        ("十、費用與補償",
         "本試驗不額外收取費用。（請說明受試者參與試驗之交通費補助或其他補償方式）"),
        ("十一、試驗相關傷害之補償",
         "若您因參與本試驗而受到傷害，將依相關法規提供適當之醫療照護與補償。"
         "（請說明試驗相關傷害之醫療處理及補償機制）"),
        ("十二、保密性",
         "您的個人資料將依法受到保護。研究人員將採取適當之保密措施，"
         "確保您的個人隱私不會被洩漏。研究結果發表時，將不會揭露您的身分。"
         "但依法規要求，主管機關、人體試驗委員會及試驗委託者得於必要時查閱您的研究資料。"),
        ("十三、自願參加與退出",
         "您的參與完全自願。您可以在任何時候退出試驗，不會影響您的醫療照護品質。"
         "您的退出不需要任何理由，也不會影響您與醫療人員之間的關係。"),
    ]

    for title, content in sections:
        add_p(doc, title, bold=True, size=12, sa=Pt(8), sb=Pt(2))
        add_p(doc, content, size=12, sa=Pt(2), sb=Pt(2))

    _consent_contact_section(doc, config, "十四")
    _consent_rights_section(doc, "十五")
    _consent_signature_block(doc)

    add_footer(doc, "2", "IRB.SF092", "2023/11/01")
    return _save(doc, output_dir, "SF092_受試者同意書（醫療器材）.docx")


# ---------------------------------------------------------------------------
# ALL_GENERATORS registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF003": generate_sf003,
    "SF004": generate_sf004,
    "SF005": generate_sf005,
    "SF062": generate_sf062,
    "SF063": generate_sf063,
    "SF075": generate_sf075,
    "SF090": generate_sf090,
    "SF091": generate_sf091,
    "SF092": generate_sf092,
}
