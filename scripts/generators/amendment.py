"""Amendment review form generators for IRB amendment review (修正案審查) at KFSYSCC.

Generates SF014, SF015, SF016 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# SF014 — 修正案審查送審資料表 (v10, 2025.03.03)
# ---------------------------------------------------------------------------


def generate_sf014(config, output_dir):
    """Generate SF014: amendment review submission checklist."""
    doc = init_doc(sz=12)
    is_drug = config["study"].get("drug_device", False)
    affects_consent = config["amendment"].get("affects_consent", False)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "修正案審查送審資料表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Admin line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿",
          size=12, sa=Pt(8))

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        (True, "※", "送審文件電子檔", "email至irb@kfsyscc.org"),
        (True, "※", "送審資料表及文件繳交完成簽收表", ""),
        (True, "※", "修正案申請表（主持人簽名及日期）", "IRB.SF015"),
        (True, "※", "修正前後對照表", "IRB.SF016"),
        (True, "※", "修正後研究計畫書（以劃線標註修改處）", ""),
        (affects_consent, "", "修正後受試者同意書（以劃線標註修改處）", ""),
        (True, "※", "顯著財務利益申報表", "IRB.SF094"),
        (is_drug, "", "臨床試驗/研究許可證明", "IRB.SF011"),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (required, cat_mark, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], check(required), size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], cat_mark, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=10)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "10", "IRB.SF014", "2025/03/03")

    path = os.path.join(output_dir, "SF014_修正案審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF015 — 修正案申請表 (v6, 2022.09.01)
# ---------------------------------------------------------------------------


def generate_sf015(config, output_dir):
    """Generate SF015: amendment application form."""
    doc = init_doc(sz=12)
    study = config["study"]
    pi = config["pi"]
    amendment = config["amendment"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "修正案申請表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、修正案基本資料 ----
    add_p(doc, "一、修正案基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=5, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "IRB編號", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], study["irb_no"], size=10)
    add_ct(tbl1.rows[1].cells[0], "計畫名稱", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1],
           f"{study['title_zh']}\n{study['title_en']}", size=10)
    add_ct(tbl1.rows[2].cells[0], "主持人", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], pi["name"], size=10)
    add_ct(tbl1.rows[3].cells[0], "單位/職稱", bold=True, size=10)
    add_ct(tbl1.rows[3].cells[1], pi["dept"], size=10)
    add_ct(tbl1.rows[4].cells[0], "修正案編號", bold=True, size=10)
    add_ct(tbl1.rows[4].cells[1], "第＿次修正", size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、修正原因 ----
    add_p(doc, "二、修正原因", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□ 研究設計變更　□ 收案條件變更　□ 受試者同意書變更",
          size=10, sa=Pt(2))
    add_p(doc, "□ 研究人員變更　□ 其他（請說明）",
          size=10, sa=Pt(4))

    desc = amendment.get("change_description", "（請說明修正內容）")
    add_p(doc, f"修正說明：{desc}", size=10, sa=Pt(8))
    doc.add_paragraph()

    # ---- 三、修正內容概述 ----
    add_p(doc, "三、修正內容概述", bold=True, size=12, sa=Pt(6))

    affects_risk = amendment.get("affects_risk", False)
    affects_consent = amendment.get("affects_consent", False)

    tbl3 = doc.add_table(rows=2, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl3.rows[0].cells[0], "是否影響受試者風險", bold=True, size=10)
    add_ct(tbl3.rows[0].cells[1],
           f"{check(affects_risk)}是　{check(not affects_risk)}否", size=10)
    add_ct(tbl3.rows[1].cells[0], "是否影響知情同意", bold=True, size=10)
    add_ct(tbl3.rows[1].cells[1],
           f"{check(affects_consent)}是　{check(not affects_consent)}否", size=10)
    apply_tb(tbl3)
    doc.add_paragraph()

    # ---- 四、主持人聲明 ----
    add_p(doc, "四、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明以上所填資料均屬正確，修正內容已詳實說明。",
          size=10, sa=Pt(4))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "6", "IRB.SF015", "2022/09/01")

    path = os.path.join(output_dir, "SF015_修正案申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF016 — 修正前後對照表 (v1, 2013.02.20)
# ---------------------------------------------------------------------------


def generate_sf016(config, output_dir):
    """Generate SF016: pre/post amendment comparison table."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "修正前後對照表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Instruction
    add_p(doc, "請以表格方式列出修正前後之對照內容", size=12, sa=Pt(8))

    # Comparison table: 4 columns, header + 5 empty rows
    cols = ["項次", "修正項目", "修正前內容", "修正後內容"]
    num_rows = 5
    tbl = doc.add_table(rows=1 + num_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Empty rows with row numbers
    for ri in range(1, num_rows + 1):
        add_ct(tbl.rows[ri].cells[0], str(ri), size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], "", size=10)
        add_ct(tbl.rows[ri].cells[2], "", size=10)
        add_ct(tbl.rows[ri].cells[3], "", size=10)

    apply_tb(tbl)
    doc.add_paragraph()

    # Note
    add_p(doc, "請以劃線方式標示修正處", size=10, sa=Pt(8))

    # Signature block
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "1", "IRB.SF016", "2013/02/20")

    path = os.path.join(output_dir, "SF016_修正前後對照表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF014": generate_sf014,
    "SF015": generate_sf015,
    "SF016": generate_sf016,
}
