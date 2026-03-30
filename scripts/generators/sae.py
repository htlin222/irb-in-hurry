"""SAE and non-compliance event generators for KFSYSCC IRB.

Generates SF079, SF044, SF074, SF080, SF024 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# SF079 — 嚴重不良反應事件審查送審資料表 (v4)
# ---------------------------------------------------------------------------


def generate_sf079(config, output_dir):
    """Generate SF079: SAE review submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "嚴重不良反應事件審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿", size=11)

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", ""),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "嚴重不良反應事件通報表（本院）", "IRB.SF044"),
        ("□", "", "嚴重不良反應事件通報表（他院）", "IRB.SF074"),
        ("□", "", "相關佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "4", "IRB.SF079", "2025/03/03")

    path = os.path.join(output_dir, "SF079_嚴重不良反應事件審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF044 — 嚴重不良反應事件通報表（本院）(v3)
# ---------------------------------------------------------------------------


def generate_sf044(config, output_dir):
    """Generate SF044: SAE report for our institution."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "嚴重不良反應事件通報表（本院）", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、事件基本資料 ----
    add_p(doc, "一、事件基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=4, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_ct(tbl1.rows[0].cells[0], "受試者編號", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], "", size=10)

    add_ct(tbl1.rows[1].cells[0], "事件發生日期", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1], "", size=10)

    add_ct(tbl1.rows[2].cells[0], "通報日期", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], "", size=10)

    add_ct(tbl1.rows[3].cells[0], "事件類別", bold=True, size=10)
    add_ct(tbl1.rows[3].cells[1],
           "□死亡　□危及生命　□永久性殘疾\n"
           "□住院或延長住院　□先天性異常\n"
           "□其他重要醫學事件", size=10)

    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、事件描述 ----
    add_p(doc, "二、事件描述", bold=True, size=12, sa=Pt(6))
    add_p(doc, "請詳述事件經過、處置及結果：", size=10, sa=Pt(4))

    # Large blank area
    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.rows[0].height = Pt(200)
    add_ct(tbl2.rows[0].cells[0], "", size=10)
    apply_tb(tbl2)
    doc.add_paragraph()

    # ---- 三、因果關係評估 ----
    add_p(doc, "三、因果關係評估", bold=True, size=12, sa=Pt(6))
    add_p(doc, "與試驗藥品／處置之關係：", size=10, sa=Pt(4))
    add_p(doc, "□確定有關　□極有可能有關　□可能有關\n"
               "□可能無關　□確定無關　□無法評估", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 四、事件結果 ----
    add_p(doc, "四、事件結果", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□已恢復　□恢復中　□未恢復　□死亡　□不詳",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、後續處理措施 ----
    add_p(doc, "五、後續處理措施", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□繼續試驗，不做修改　□繼續試驗，修改劑量／方案\n"
               "□暫停試驗　□終止試驗　□其他", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 六、通報人資訊 ----
    add_p(doc, "六、通報人資訊", bold=True, size=12, sa=Pt(6))

    tbl6 = doc.add_table(rows=3, cols=2)
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl6.rows[0].cells[0], "計畫主持人", bold=True, size=10)
    add_ct(tbl6.rows[0].cells[1], pi["name"], size=10)
    add_ct(tbl6.rows[1].cells[0], "聯絡電話", bold=True, size=10)
    add_ct(tbl6.rows[1].cells[1], pi.get("phone", ""), size=10)
    add_ct(tbl6.rows[2].cells[0], "電子信箱", bold=True, size=10)
    add_ct(tbl6.rows[2].cells[1], pi.get("email", ""), size=10)
    apply_tb(tbl6)
    doc.add_paragraph()

    # Signature block
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "3", "IRB.SF044", "2022/09/01")

    path = os.path.join(output_dir, "SF044_嚴重不良反應事件通報表_本院.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF074 — 嚴重不良反應事件通報表（他院）(v3)
# ---------------------------------------------------------------------------


def generate_sf074(config, output_dir):
    """Generate SF074: SAE report from other institutions."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "嚴重不良反應事件通報表（他院）", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、事件基本資料 ----
    add_p(doc, "一、事件基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=6, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_ct(tbl1.rows[0].cells[0], "通報來源機構名稱", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], "", size=10)

    add_ct(tbl1.rows[1].cells[0], "原始通報日期", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1], "", size=10)

    add_ct(tbl1.rows[2].cells[0], "受試者編號", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], "", size=10)

    add_ct(tbl1.rows[3].cells[0], "事件發生日期", bold=True, size=10)
    add_ct(tbl1.rows[3].cells[1], "", size=10)

    add_ct(tbl1.rows[4].cells[0], "通報日期", bold=True, size=10)
    add_ct(tbl1.rows[4].cells[1], "", size=10)

    add_ct(tbl1.rows[5].cells[0], "事件類別", bold=True, size=10)
    add_ct(tbl1.rows[5].cells[1],
           "□死亡　□危及生命　□永久性殘疾\n"
           "□住院或延長住院　□先天性異常\n"
           "□其他重要醫學事件", size=10)

    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、事件描述 ----
    add_p(doc, "二、事件描述", bold=True, size=12, sa=Pt(6))
    add_p(doc, "請詳述事件經過、處置及結果：", size=10, sa=Pt(4))

    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.rows[0].height = Pt(200)
    add_ct(tbl2.rows[0].cells[0], "", size=10)
    apply_tb(tbl2)
    doc.add_paragraph()

    # ---- 三、因果關係評估 ----
    add_p(doc, "三、因果關係評估", bold=True, size=12, sa=Pt(6))
    add_p(doc, "與試驗藥品／處置之關係：", size=10, sa=Pt(4))
    add_p(doc, "□確定有關　□極有可能有關　□可能有關\n"
               "□可能無關　□確定無關　□無法評估", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 四、事件結果 ----
    add_p(doc, "四、事件結果", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□已恢復　□恢復中　□未恢復　□死亡　□不詳",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、後續處理措施 ----
    add_p(doc, "五、後續處理措施", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□繼續試驗，不做修改　□繼續試驗，修改劑量／方案\n"
               "□暫停試驗　□終止試驗　□其他", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 六、通報人資訊 ----
    add_p(doc, "六、通報人資訊", bold=True, size=12, sa=Pt(6))

    tbl6 = doc.add_table(rows=3, cols=2)
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl6.rows[0].cells[0], "計畫主持人", bold=True, size=10)
    add_ct(tbl6.rows[0].cells[1], pi["name"], size=10)
    add_ct(tbl6.rows[1].cells[0], "聯絡電話", bold=True, size=10)
    add_ct(tbl6.rows[1].cells[1], pi.get("phone", ""), size=10)
    add_ct(tbl6.rows[2].cells[0], "電子信箱", bold=True, size=10)
    add_ct(tbl6.rows[2].cells[1], pi.get("email", ""), size=10)
    apply_tb(tbl6)
    doc.add_paragraph()

    # Signature block
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "3", "IRB.SF074", "2022/09/01")

    path = os.path.join(output_dir, "SF074_嚴重不良反應事件通報表_他院.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF080 — 試驗不遵從事件審查送審資料表 (v4)
# ---------------------------------------------------------------------------


def generate_sf080(config, output_dir):
    """Generate SF080: non-compliance submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "試驗不遵從事件審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿", size=11)

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", ""),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "試驗不遵從事件報告表", "IRB.SF024"),
        ("□", "", "相關佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "4", "IRB.SF080", "2025/03/03")

    path = os.path.join(output_dir, "SF080_試驗不遵從事件審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF024 — 試驗不遵從事件報告表 (v6)
# ---------------------------------------------------------------------------


def generate_sf024(config, output_dir):
    """Generate SF024: non-compliance / protocol deviation report."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "試驗不遵從事件報告表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、事件基本資料 ----
    add_p(doc, "一、事件基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=4, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_ct(tbl1.rows[0].cells[0], "事件發生日期", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], "", size=10)

    add_ct(tbl1.rows[1].cells[0], "發現日期", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1], "", size=10)

    add_ct(tbl1.rows[2].cells[0], "通報日期", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], "", size=10)

    add_ct(tbl1.rows[3].cells[0], "事件類別", bold=True, size=10)
    add_ct(tbl1.rows[3].cells[1],
           "□計畫書偏離　□未遵醫囑　□違規收案\n"
           "□同意書問題　□其他", size=10)

    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、事件描述 ----
    add_p(doc, "二、事件描述", bold=True, size=12, sa=Pt(6))
    add_p(doc, "請詳述事件經過：", size=10, sa=Pt(4))

    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.rows[0].height = Pt(200)
    add_ct(tbl2.rows[0].cells[0], "", size=10)
    apply_tb(tbl2)
    doc.add_paragraph()

    # ---- 三、事件分類 ----
    add_p(doc, "三、事件分類", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□ 重大偏離（可能影響受試者安全或研究完整性）", size=10, sa=Pt(4))
    add_p(doc, "□ 輕微偏離（不影響受試者安全或研究完整性）", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 四、矯正及預防措施 (CAPA) ----
    add_p(doc, "四、矯正及預防措施（CAPA）", bold=True, size=12, sa=Pt(6))

    add_p(doc, "已採取或預計採取之矯正措施：", size=10, sa=Pt(4))
    tbl4a = doc.add_table(rows=1, cols=1)
    tbl4a.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4a.rows[0].height = Pt(120)
    add_ct(tbl4a.rows[0].cells[0], "", size=10)
    apply_tb(tbl4a)
    doc.add_paragraph()

    add_p(doc, "預防再次發生之措施：", size=10, sa=Pt(4))
    tbl4b = doc.add_table(rows=1, cols=1)
    tbl4b.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4b.rows[0].height = Pt(120)
    add_ct(tbl4b.rows[0].cells[0], "", size=10)
    apply_tb(tbl4b)
    doc.add_paragraph()

    # ---- 五、通報人資訊 ----
    add_p(doc, "五、通報人資訊", bold=True, size=12, sa=Pt(6))

    tbl5 = doc.add_table(rows=3, cols=2)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl5.rows[0].cells[0], "計畫主持人", bold=True, size=10)
    add_ct(tbl5.rows[0].cells[1], pi["name"], size=10)
    add_ct(tbl5.rows[1].cells[0], "聯絡電話", bold=True, size=10)
    add_ct(tbl5.rows[1].cells[1], pi.get("phone", ""), size=10)
    add_ct(tbl5.rows[2].cells[0], "電子信箱", bold=True, size=10)
    add_ct(tbl5.rows[2].cells[1], pi.get("email", ""), size=10)
    apply_tb(tbl5)
    doc.add_paragraph()

    # Signature block
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "6", "IRB.SF024", "2025/03/03")

    path = os.path.join(output_dir, "SF024_試驗不遵從事件報告表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF079": generate_sf079,
    "SF044": generate_sf044,
    "SF074": generate_sf074,
    "SF080": generate_sf080,
    "SF024": generate_sf024,
}
