"""Suspension/early termination form generators for KFSYSCC IRB.

Generates SF047, SF048 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# SF047 — 計畫暫停/提前終止審查送審資料表 (v4)
# ---------------------------------------------------------------------------


def generate_sf047(config, output_dir):
    """Generate SF047: suspension/early termination submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "計畫暫停/提前終止審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿", size=11)

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", ""),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "計畫暫停/提前終止報告書（主持人簽名）", "IRB.SF048"),
        ("□", "", "相關佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Note
    add_p(doc, "一式二份（正本一份、影本一份）", size=10, sa=Pt(4))

    # Footer
    add_footer(doc, "4", "IRB.SF047", "2022/09/01")

    path = os.path.join(output_dir, "SF047_計畫暫停提前終止審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF048 — 計畫暫停/提前終止報告書 (v4)
# ---------------------------------------------------------------------------


def generate_sf048(config, output_dir):
    """Generate SF048: suspension/early termination report."""
    doc = init_doc(sz=12)
    pi = config["pi"]
    subjects = config["subjects"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "計畫暫停/提前終止報告書", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、計畫狀態 ----
    add_p(doc, "一、計畫狀態", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□暫停　□提前終止", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 二、暫停/終止日期 ----
    add_p(doc, "二、暫停/終止日期", bold=True, size=12, sa=Pt(6))
    add_p(doc, "＿＿＿＿年＿＿月＿＿日", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 三、暫停/終止原因 ----
    add_p(doc, "三、暫停/終止原因", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□安全性考量　□療效不佳　□收案困難\n"
               "□經費不足　□主持人因素　□主管機關要求\n"
               "□其他：＿＿＿＿", size=10, sa=Pt(4))

    add_p(doc, "詳細說明：", size=10, sa=Pt(4))
    tbl3 = doc.add_table(rows=1, cols=1)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.rows[0].height = Pt(150)
    add_ct(tbl3.rows[0].cells[0], "", size=10)
    apply_tb(tbl3)
    doc.add_paragraph()

    # ---- 四、受試者處置 ----
    add_p(doc, "四、受試者處置", bold=True, size=12, sa=Pt(6))

    tbl4 = doc.add_table(rows=2, cols=2)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl4.rows[0].cells[0], "預計收案人數", bold=True, size=10)
    add_ct(tbl4.rows[0].cells[1], str(subjects.get("planned_n", 0)), size=10)
    add_ct(tbl4.rows[1].cells[0], "實際收案人數", bold=True, size=10)
    add_ct(tbl4.rows[1].cells[1], str(subjects.get("actual_n", 0)), size=10)
    apply_tb(tbl4)
    doc.add_paragraph()

    add_p(doc, "受試者後續處理方式：", size=10, sa=Pt(4))
    tbl4b = doc.add_table(rows=1, cols=1)
    tbl4b.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4b.rows[0].height = Pt(120)
    add_ct(tbl4b.rows[0].cells[0], "", size=10)
    apply_tb(tbl4b)
    doc.add_paragraph()

    # ---- 五、藥品/檢體處理 ----
    add_p(doc, "五、藥品/檢體處理", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□已銷毀　□退還　□繼續保存　□不適用", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 六、主持人聲明 ----
    add_p(doc, "六、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明上述資料均屬正確，並已妥善處理受試者相關事宜。",
          size=10, sa=Pt(4))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "4", "IRB.SF048", "2022/09/01")

    path = os.path.join(output_dir, "SF048_計畫暫停提前終止報告書.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF047": generate_sf047,
    "SF048": generate_sf048,
}
