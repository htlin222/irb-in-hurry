"""Appeal form generators for KFSYSCC IRB.

Generates SF077, SF054 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# SF077 — 申覆案審查送審資料表 (v5)
# ---------------------------------------------------------------------------


def generate_sf077(config, output_dir):
    """Generate SF077: appeal review submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "申覆案審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿", size=11)

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", ""),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "申覆案申請表（主持人簽名）", "IRB.SF054"),
        ("■", "※", "原審查結果通知書影本", ""),
        ("□", "", "相關佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Note
    add_p(doc, "研究人員收到否決通知一個月內得提出申覆", size=10, sa=Pt(4))

    # Footer
    add_footer(doc, "5", "IRB.SF077", "2022/09/01")

    path = os.path.join(output_dir, "SF077_申覆案審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF054 — 申覆案申請表 (v3)
# ---------------------------------------------------------------------------


def generate_sf054(config, output_dir):
    """Generate SF054: appeal application form."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "申覆案申請表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、原審查案基本資料 ----
    add_p(doc, "一、原審查案基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=5, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "IRB編號", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], config["study"]["irb_no"], size=10)
    add_ct(tbl1.rows[1].cells[0], "計畫名稱", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1], config["study"]["title_zh"], size=10)
    add_ct(tbl1.rows[2].cells[0], "主持人", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], pi["name"], size=10)
    add_ct(tbl1.rows[3].cells[0], "原審查結果", bold=True, size=10)
    add_ct(tbl1.rows[3].cells[1], "", size=10)
    add_ct(tbl1.rows[4].cells[0], "原審查日期", bold=True, size=10)
    add_ct(tbl1.rows[4].cells[1], "", size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、申覆理由 ----
    add_p(doc, "二、申覆理由", bold=True, size=12, sa=Pt(6))
    add_p(doc, "請逐項說明對審查意見之回覆：", size=10, sa=Pt(4))

    tbl2 = doc.add_table(rows=1, cols=1)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.rows[0].height = Pt(200)
    add_ct(tbl2.rows[0].cells[0],
           "1.\n\n2.\n\n3.", size=10)
    apply_tb(tbl2)
    doc.add_paragraph()

    # ---- 三、修正說明 ----
    add_p(doc, "三、修正說明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "請說明針對審查意見所做之修正：", size=10, sa=Pt(4))

    tbl3 = doc.add_table(rows=1, cols=1)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.rows[0].height = Pt(150)
    add_ct(tbl3.rows[0].cells[0], "", size=10)
    apply_tb(tbl3)
    doc.add_paragraph()

    # ---- 四、附件說明 ----
    add_p(doc, "四、附件說明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□修正後計畫書　□修正後同意書　□其他（請說明）：＿＿＿＿",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、主持人聲明 ----
    add_p(doc, "五、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明上述資料均屬正確，並已依審查意見進行修正。",
          size=10, sa=Pt(4))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "3", "IRB.SF054", "2022/09/01")

    path = os.path.join(output_dir, "SF054_申覆案申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF077": generate_sf077,
    "SF054": generate_sf054,
}
