"""Import review form generators for KFSYSCC IRB.

Generates SF066, SF067, SF068, SF093 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


# ---------------------------------------------------------------------------
# SF066 — 專案進口審查送審資料表 (v5)
# ---------------------------------------------------------------------------


def generate_sf066(config, output_dir):
    """Generate SF066: import review submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "專案進口審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", ""),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "專案進口申請表（主持人簽名）", "IRB.SF067"),
        ("■", "※", "專案進口受試者同意書", "IRB.SF068"),
        ("■", "※", "藥品/醫材相關資料", ""),
        ("□", "", "其他佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "5", "IRB.SF066", "2022/09/01")

    path = os.path.join(output_dir, "SF066_專案進口審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF067 — 專案進口申請表 (v4)
# ---------------------------------------------------------------------------


def generate_sf067(config, output_dir):
    """Generate SF067: import application form."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "專案進口申請表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、基本資料 ----
    add_p(doc, "一、基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=3, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "IRB編號", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], config["study"]["irb_no"], size=10)
    add_ct(tbl1.rows[1].cells[0], "計畫名稱", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1], config["study"]["title_zh"], size=10)
    add_ct(tbl1.rows[2].cells[0], "主持人", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], pi["name"], size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、進口藥品/醫材資料 ----
    add_p(doc, "二、進口藥品/醫材資料", bold=True, size=12, sa=Pt(6))

    import_cols = ["品名", "成分/規格", "數量", "用途", "製造廠", "輸入國"]
    tbl2 = doc.add_table(rows=4, cols=6)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, col_name in enumerate(import_cols):
        cell = tbl2.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")
    for ri in range(1, 4):
        for ci in range(6):
            add_ct(tbl2.rows[ri].cells[ci], "", size=10)
    apply_tb(tbl2)
    doc.add_paragraph()

    # ---- 三、使用說明 ----
    add_p(doc, "三、使用說明", bold=True, size=12, sa=Pt(6))

    tbl3 = doc.add_table(rows=3, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl3.rows[0].cells[0], "適應症", bold=True, size=10)
    add_ct(tbl3.rows[0].cells[1], "", size=10)
    add_ct(tbl3.rows[1].cells[0], "用法用量", bold=True, size=10)
    add_ct(tbl3.rows[1].cells[1], "", size=10)
    add_ct(tbl3.rows[2].cells[0], "使用期間", bold=True, size=10)
    add_ct(tbl3.rows[2].cells[1], "", size=10)
    apply_tb(tbl3)
    doc.add_paragraph()

    # ---- 四、進口理由 ----
    add_p(doc, "四、進口理由", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□國內無替代品　□臨床試驗需要　□緊急醫療需要　□其他：＿＿＿＿",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、主持人聲明 ----
    add_p(doc, "五、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明上述資料均屬正確，並將依核准之計畫書使用進口藥品/醫材。",
          size=10, sa=Pt(4))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "4", "IRB.SF067", "2022/09/01")

    path = os.path.join(output_dir, "SF067_專案進口申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF068 — 專案進口受試者同意書 (v3)
# ---------------------------------------------------------------------------


def generate_sf068(config, output_dir):
    """Generate SF068: import consent form template."""
    doc = init_doc(sz=12)
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "專案進口受試者同意書", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Study info
    add_p(doc, f"計畫名稱：{config['study']['title_zh']}", size=11, sa=Pt(4))
    add_p(doc, f"主持人：{pi['name']}", size=11, sa=Pt(8))

    # ---- 一、說明 ----
    add_p(doc, "一、說明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "（請說明進口藥品/醫療器材之名稱、成分、規格及進口原因）",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 二、使用目的 ----
    add_p(doc, "二、使用目的", bold=True, size=12, sa=Pt(6))
    add_p(doc, "（請說明該藥品/醫療器材之使用目的及預期效益）",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 三、可能風險與副作用 ----
    add_p(doc, "三、可能風險與副作用", bold=True, size=12, sa=Pt(6))
    add_p(doc, "（請說明可能產生之風險、副作用及其發生機率）",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 四、替代方案 ----
    add_p(doc, "四、替代方案", bold=True, size=12, sa=Pt(6))
    add_p(doc, "（請說明除使用進口藥品/醫療器材外，是否有其他替代治療方式）",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、受試者權益 ----
    add_p(doc, "五、受試者權益", bold=True, size=12, sa=Pt(6))
    add_p(doc, "1. 您有權利隨時退出本計畫，且不會影響您應得的醫療照護。",
          size=10, sa=Pt(2))
    add_p(doc, "2. 您的個人資料將依法予以保密。",
          size=10, sa=Pt(2))
    add_p(doc, "3. 如有任何與本計畫相關之傷害，將依相關法規給予適當補償。",
          size=10, sa=Pt(2))
    add_p(doc, "4. 如有任何疑問，可隨時與計畫主持人聯繫。",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 六、同意聲明 ----
    add_p(doc, "六、同意聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人已充分瞭解上述說明，同意使用專案進口之藥品/醫療器材。",
          size=10, sa=Pt(8))

    # Signature lines
    sig_tbl = doc.add_table(rows=4, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(sig_tbl.rows[0].cells[0], "受試者簽名：＿＿＿＿＿＿＿＿", size=10)
    add_ct(sig_tbl.rows[0].cells[1], "日期：＿＿＿＿年＿＿月＿＿日", size=10)
    add_ct(sig_tbl.rows[1].cells[0], "法定代理人簽名：＿＿＿＿＿＿", size=10)
    add_ct(sig_tbl.rows[1].cells[1], "日期：＿＿＿＿年＿＿月＿＿日", size=10)
    add_ct(sig_tbl.rows[2].cells[0], "見證人簽名：＿＿＿＿＿＿＿＿", size=10)
    add_ct(sig_tbl.rows[2].cells[1], "日期：＿＿＿＿年＿＿月＿＿日", size=10)
    add_ct(sig_tbl.rows[3].cells[0], "主持人簽名：＿＿＿＿＿＿＿＿", size=10)
    add_ct(sig_tbl.rows[3].cells[1], "日期：＿＿＿＿年＿＿月＿＿日", size=10)
    apply_tb(sig_tbl)

    doc.add_paragraph()

    # Footer
    add_footer(doc, "3", "IRB.SF068", "2018/01/02")

    path = os.path.join(output_dir, "SF068_專案進口受試者同意書.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF093 — 專案進口許可證明 (v1)
# ---------------------------------------------------------------------------


def generate_sf093(config, output_dir):
    """Generate SF093: import permit certificate."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "專案進口許可證明", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Certificate statement
    add_p(doc, "茲證明下列藥品/醫療器材經本院人體試驗委員會審查核准專案進口使用：",
          size=11, sa=Pt(8))

    # Info table
    info_cols = ["品名", "成分", "數量", "進口國", "核准日期", "有效期限"]
    tbl = doc.add_table(rows=4, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, col_name in enumerate(info_cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")
    for ri in range(1, 4):
        for ci in range(6):
            add_ct(tbl.rows[ri].cells[ci], "", size=10)
    apply_tb(tbl)
    doc.add_paragraph()

    # Approval number
    add_p(doc, "IRB審查核准文號：＿＿＿＿＿＿＿＿＿＿", size=11, sa=Pt(8))

    # Chairman signature
    add_p(doc, "人體試驗委員會 主任委員：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿年＿＿月＿＿日",
          size=11, sa=Pt(12))

    # Footer
    add_footer(doc, "1", "IRB.SF093", "2022/09/01")

    path = os.path.join(output_dir, "SF093_專案進口許可證明.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF066": generate_sf066,
    "SF067": generate_sf067,
    "SF068": generate_sf068,
    "SF093": generate_sf093,
}
