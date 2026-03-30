"""IB update and multi-center letter generators for KFSYSCC IRB.

Generates SF082, SF083, SF084, SF085 forms from config dict.
SF082/SF083: 更新主持人手冊 (Investigator's Brochure update)
SF084/SF085: 多中心信函 (multi-center letter)
"""
import os

from scripts.docx_utils import *
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# SF082 — 更新主持人手冊審查送審資料表 (v5, 2025.03.03)
# ---------------------------------------------------------------------------


def generate_sf082(config, output_dir):
    """Generate SF082: IB update submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "更新主持人手冊審查送審資料表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿",
          size=12, sa=Pt(8))

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", "email至irb@kfsyscc.org"),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "更新主持人手冊申請表（主持人簽名）", "IRB.SF083"),
        ("■", "※", "更新之主持人手冊", ""),
        ("■", "※", "更新內容摘要/變更說明", ""),
        ("□", "", "相關佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "5", "IRB.SF082", "2025/03/03")

    path = os.path.join(output_dir, "SF082_更新主持人手冊審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF083 — 更新主持人手冊申請表 (v2, 2018.01.02)
# ---------------------------------------------------------------------------


def generate_sf083(config, output_dir):
    """Generate SF083: IB update application form."""
    doc = init_doc(sz=12)
    study = config["study"]
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "更新主持人手冊申請表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、基本資料 ----
    add_p(doc, "一、基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=3, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "IRB編號", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], study["irb_no"], size=10)
    add_ct(tbl1.rows[1].cells[0], "計畫名稱", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1],
           f"{study['title_zh']}\n{study['title_en']}", size=10)
    add_ct(tbl1.rows[2].cells[0], "主持人", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], pi["name"], size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、更新內容 ----
    add_p(doc, "二、更新內容", bold=True, size=12, sa=Pt(6))

    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl2.rows[0].cells[0], "主持人手冊版本", bold=True, size=10)
    add_ct(tbl2.rows[0].cells[1], "第＿版，日期：＿＿＿＿", size=10)
    add_ct(tbl2.rows[1].cells[0], "更新重點摘要", bold=True, size=10)
    add_ct(tbl2.rows[1].cells[1], "（請說明更新重點內容）", size=10)
    apply_tb(tbl2)

    # Blank area for user to fill in details
    tbl2b = doc.add_table(rows=1, cols=1)
    tbl2b.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2b.rows[0].height = Pt(150)
    add_ct(tbl2b.rows[0].cells[0], "", size=10)
    apply_tb(tbl2b)
    doc.add_paragraph()

    # ---- 三、風險評估 ----
    add_p(doc, "三、風險評估", bold=True, size=12, sa=Pt(6))

    tbl3 = doc.add_table(rows=3, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl3.rows[0].cells[0], "更新內容是否改變受試者之風險", bold=True, size=10)
    add_ct(tbl3.rows[0].cells[1], "□是　□否", size=10)
    add_ct(tbl3.rows[1].cells[0], "是否需修正受試者同意書", bold=True, size=10)
    add_ct(tbl3.rows[1].cells[1], "□是　□否", size=10)
    add_ct(tbl3.rows[2].cells[0], "是否需修正研究計畫書", bold=True, size=10)
    add_ct(tbl3.rows[2].cells[1], "□是　□否", size=10)
    apply_tb(tbl3)
    doc.add_paragraph()

    # ---- 四、主持人聲明 ----
    add_p(doc, "四、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明以上所填資料均屬正確，更新內容已詳實說明。",
          size=10, sa=Pt(4))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "2", "IRB.SF083", "2018/01/02")

    path = os.path.join(output_dir, "SF083_更新主持人手冊申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF084 — 多中心信函審查送審資料表 (v3, 2022.09.01)
# ---------------------------------------------------------------------------


def generate_sf084(config, output_dir):
    """Generate SF084: multi-center letter submission checklist."""
    doc = init_doc(sz=12)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "多中心信函審查送審資料表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿",
          size=12, sa=Pt(8))

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("■", "※", "送審文件電子檔", "email至irb@kfsyscc.org"),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "多中心/通知函申請表（主持人簽名）", "IRB.SF085"),
        ("■", "※", "多中心信函/通知函原件", ""),
        ("□", "", "相關佐證資料", ""),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (ready, category, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=9)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "3", "IRB.SF084", "2022/09/01")

    path = os.path.join(output_dir, "SF084_多中心信函審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF085 — 多中心信函申請表 (v3, 2022.09.01)
# ---------------------------------------------------------------------------


def generate_sf085(config, output_dir):
    """Generate SF085: multi-center letter application form."""
    doc = init_doc(sz=12)
    study = config["study"]
    pi = config["pi"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "多中心信函申請表", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # ---- 一、基本資料 ----
    add_p(doc, "一、基本資料", bold=True, size=12, sa=Pt(6))

    tbl1 = doc.add_table(rows=3, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "IRB編號", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], study["irb_no"], size=10)
    add_ct(tbl1.rows[1].cells[0], "計畫名稱", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1],
           f"{study['title_zh']}\n{study['title_en']}", size=10)
    add_ct(tbl1.rows[2].cells[0], "主持人", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1], pi["name"], size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、信函內容摘要 ----
    add_p(doc, "二、信函內容摘要", bold=True, size=12, sa=Pt(6))

    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl2.rows[0].cells[0], "信函來源", bold=True, size=10)
    add_ct(tbl2.rows[0].cells[1], "＿＿＿＿（機構名稱）", size=10)
    add_ct(tbl2.rows[1].cells[0], "信函日期", bold=True, size=10)
    add_ct(tbl2.rows[1].cells[1], "＿＿＿＿年＿＿月＿＿日", size=10)
    add_ct(tbl2.rows[2].cells[0], "信函類型", bold=True, size=10)
    add_ct(tbl2.rows[2].cells[1],
           "□安全性報告　□收案進度　□計畫修正通知　□其他", size=10)
    add_ct(tbl2.rows[3].cells[0], "內容摘要", bold=True, size=10)
    add_ct(tbl2.rows[3].cells[1], "（請說明信函主要內容）", size=10)
    apply_tb(tbl2)

    # Blank area for content summary
    tbl2b = doc.add_table(rows=1, cols=1)
    tbl2b.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2b.rows[0].height = Pt(150)
    add_ct(tbl2b.rows[0].cells[0], "", size=10)
    apply_tb(tbl2b)
    doc.add_paragraph()

    # ---- 三、影響評估 ----
    add_p(doc, "三、影響評估", bold=True, size=12, sa=Pt(6))

    tbl3 = doc.add_table(rows=2, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl3.rows[0].cells[0], "是否影響本院研究執行", bold=True, size=10)
    add_ct(tbl3.rows[0].cells[1], "□是　□否", size=10)
    add_ct(tbl3.rows[1].cells[0], "是否需修正計畫書/同意書", bold=True, size=10)
    add_ct(tbl3.rows[1].cells[1], "□是　□否", size=10)
    apply_tb(tbl3)
    doc.add_paragraph()

    # ---- 四、主持人聲明 ----
    add_p(doc, "四、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明以上所填資料均屬正確，信函內容已詳實說明。",
          size=10, sa=Pt(4))
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "3", "IRB.SF085", "2022/09/01")

    path = os.path.join(output_dir, "SF085_多中心信函申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF082": generate_sf082,
    "SF083": generate_sf083,
    "SF084": generate_sf084,
    "SF085": generate_sf085,
}
