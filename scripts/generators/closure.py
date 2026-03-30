"""Closure form generators for IRB close-out review at KFSYSCC.

Generates SF036, SF037, SF038, SF023 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# SF036 — 結案審查送審資料表
# ---------------------------------------------------------------------------


def generate_sf036(config, output_dir):
    """Generate SF036: closure review submission checklist."""
    doc = init_doc(sz=12)
    is_drug = config["study"].get("drug_device", False)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "結案審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Checklist table: 5 columns
    cols = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    items = [
        ("1", True, True, "送審文件電子檔", ""),
        ("2", True, True, "送審資料表及文件繳交完成簽收表", ""),
        ("3", True, True, "結案報告摘要表", "IRB.SF037"),
        ("4", True, True, "中文計畫摘要", ""),
        ("5", True, True, "結案報告書", "IRB.SF038"),
        ("6", is_drug, False, "個案報告表", ""),
        ("7", is_drug, False, "最終統計報告", ""),
        ("8", is_drug, False, "藥品銷毀紀錄", ""),
        ("9", is_drug, False, "最終監測報告", ""),
        ("10", is_drug, False, "資料及安全性監測計畫報告書", "IRB.SF023"),
        ("11", not is_drug, True, "其他：資料及安全性監測計畫報告書", "IRB.SF023"),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (num, required, always, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], check(required), size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # 藥品/非藥品 column
        if always:
            cat = "通用"
        else:
            cat = "藥品" if is_drug else ""
        add_ct(tbl.rows[ri].cells[1], cat, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], f"{num}. {name}", size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=10)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "6", "IRB.SF036", "2025/03/03")

    path = os.path.join(output_dir, "SF036_結案審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF037 — 結案報告摘要表
# ---------------------------------------------------------------------------


def generate_sf037(config, output_dir):
    """Generate SF037: closure report summary form."""
    doc = init_doc(sz=12)
    study = config["study"]
    pi = config["pi"]
    dates = config["dates"]
    subjects = config["subjects"]
    closure = config["closure"]
    data_safety = closure["data_safety"]
    is_retro = study.get("type", "") == "retrospective"

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "結案報告摘要表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header
    add_header(doc, config)

    # ---- 一、計畫期間 ----
    add_p(doc, "一、計畫期間", bold=True, size=12, sa=Pt(6))

    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    items_1 = [
        ("核准起迄期間", f"{dates.get('study_start', '')} 至 {dates.get('study_end', '')}"),
        ("展延次數", str(closure.get("extensions", 0))),
        ("修正案次數", str(closure.get("amendments", 0))),
        ("實際收案期間", dates.get("data_period", "")),
    ]
    for ri, (label, val) in enumerate(items_1):
        add_ct(tbl.rows[ri].cells[0], label, bold=True, size=10)
        add_ct(tbl.rows[ri].cells[1], val, size=10)
    apply_tb(tbl)
    doc.add_paragraph()

    # ---- 二、核准條件 ----
    add_p(doc, "二、核准條件", bold=True, size=12, sa=Pt(6))
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl2.rows[0].cells[0], "資料及安全性監測計畫", bold=True, size=10)
    add_ct(tbl2.rows[0].cells[1], f"{check(True)}有  {check(False)}無", size=10)
    add_ct(tbl2.rows[1].cells[0], "受試者同意書", bold=True, size=10)
    waiver = subjects.get("consent_waiver", False)
    add_ct(tbl2.rows[1].cells[1],
           f"{check(not waiver)}有  {check(waiver)}免除（免除同意書）", size=10)
    apply_tb(tbl2)
    doc.add_paragraph()

    # ---- 三、收錄個案描述 ----
    add_p(doc, "三、收錄個案描述", bold=True, size=12, sa=Pt(6))
    tbl3 = doc.add_table(rows=2, cols=2)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl3.rows[0].cells[0], "預計收案人數", bold=True, size=10)
    add_ct(tbl3.rows[0].cells[1], str(subjects.get("planned_n", 0)), size=10)
    add_ct(tbl3.rows[1].cells[0], "實際收案人數", bold=True, size=10)
    add_ct(tbl3.rows[1].cells[1], str(subjects.get("actual_n", 0)), size=10)
    apply_tb(tbl3)

    # Groups sub-table
    groups = subjects.get("groups", [])
    if groups:
        doc.add_paragraph()
        tbl_g = doc.add_table(rows=1 + len(groups), cols=2)
        tbl_g.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_ct(tbl_g.rows[0].cells[0], "組別", bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl_g.rows[0].cells[1], "人數", bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(tbl_g.rows[0].cells[0], "D9E2F3")
        set_cell_shading(tbl_g.rows[0].cells[1], "D9E2F3")
        for gi, grp in enumerate(groups, start=1):
            add_ct(tbl_g.rows[gi].cells[0], grp.get("name", ""), size=10)
            add_ct(tbl_g.rows[gi].cells[1], str(grp.get("n", 0)), size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
        apply_tb(tbl_g)

    add_p(doc, "性別：□男 ___人  □女 ___人", size=10, sa=Pt(4))
    add_p(doc, "年齡：___歲 至 ___歲（中位數：___歲）", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 四、嚴重不良反應事件 ----
    add_p(doc, "四、嚴重不良反應事件", bold=True, size=12, sa=Pt(6))
    sae_count = closure.get("sae_count", 0)
    add_p(doc, f"{check(is_retro)}不適用（回溯性研究）", size=10, sa=Pt(2))
    add_p(doc, f"{check(not is_retro and sae_count == 0)}無嚴重不良反應事件", size=10, sa=Pt(2))
    add_p(doc, f"{check(not is_retro and sae_count > 0)}有，共 {sae_count} 件（詳見第十節）",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、資料安全維護 ----
    add_p(doc, "五、資料安全維護", bold=True, size=12, sa=Pt(6))
    tbl5 = doc.add_table(rows=4, cols=2)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    items_5 = [
        ("資料去識別化", f"{check(data_safety.get('deidentified', False))}是  "
                     f"{check(not data_safety.get('deidentified', False))}否"),
        ("資料加密保存", f"{check(data_safety.get('encrypted', False))}是  "
                     f"{check(not data_safety.get('encrypted', False))}否"),
        ("資料保存年限", f"{data_safety.get('retention_years', 7)} 年"),
        ("資料存取授權人員", data_safety.get("authorized_personnel", "")),
    ]
    for ri, (label, val) in enumerate(items_5):
        add_ct(tbl5.rows[ri].cells[0], label, bold=True, size=10)
        add_ct(tbl5.rows[ri].cells[1], val, size=10)
    apply_tb(tbl5)
    doc.add_paragraph()

    # ---- 六、檢體處理 ----
    add_p(doc, "六、檢體處理", bold=True, size=12, sa=Pt(6))
    has_specimens = closure.get("specimens", False)
    add_p(doc, f"{check(not has_specimens)}本計畫無留存檢體", size=10, sa=Pt(2))
    add_p(doc, f"{check(has_specimens)}本計畫有留存檢體，處理方式如下：", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 七、計畫主持人聲明 ----
    add_p(doc, "七、計畫主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明本計畫所有研究資料均已妥善保存，並已依規定完成結案相關程序。",
          size=10, sa=Pt(4))
    add_p(doc, f"計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # ---- 八、聯絡人資訊 ----
    add_p(doc, "八、聯絡人資訊", bold=True, size=12, sa=Pt(6))
    tbl8 = doc.add_table(rows=3, cols=2)
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl8.rows[0].cells[0], "聯絡人", bold=True, size=10)
    add_ct(tbl8.rows[0].cells[1], pi["name"], size=10)
    add_ct(tbl8.rows[1].cells[0], "聯絡電話", bold=True, size=10)
    add_ct(tbl8.rows[1].cells[1], pi.get("phone", ""), size=10)
    add_ct(tbl8.rows[2].cells[0], "電子信箱", bold=True, size=10)
    add_ct(tbl8.rows[2].cells[1], pi.get("email", ""), size=10)
    apply_tb(tbl8)
    doc.add_paragraph()

    # ---- 九、收錄受試者清單 ----
    add_p(doc, "九、收錄受試者清單", bold=True, size=12, sa=Pt(6))
    add_p(doc, f"{check(waiver)}免除同意書，不適用受試者清單", size=10, sa=Pt(2))
    add_p(doc, f"{check(not waiver)}受試者清單如附件", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 十、SAE清單 ----
    add_p(doc, "十、SAE清單", bold=True, size=12, sa=Pt(6))
    add_p(doc, f"{check(is_retro)}不適用（回溯性研究）", size=10, sa=Pt(2))
    add_p(doc, f"{check(not is_retro and sae_count == 0)}無", size=10, sa=Pt(2))
    add_p(doc, f"{check(not is_retro and sae_count > 0)}有，清單如附件", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 十一、新醫療技術清單 ----
    add_p(doc, "十一、新醫療技術清單", bold=True, size=12, sa=Pt(6))
    add_p(doc, "□有  □無", size=10, sa=Pt(4))

    # Footer
    add_footer(doc, "6", "IRB.SF037", "2025/03/03")

    path = os.path.join(output_dir, "SF037_結案報告摘要表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF038 — 結案報告書
# ---------------------------------------------------------------------------


def generate_sf038(config, output_dir):
    """Generate SF038: closure report document (template structure)."""
    doc = init_doc(sz=12)
    study = config["study"]
    dates = config["dates"]

    # ---- Cover page ----
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "結案報告書", bold=True, size=18,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(24))

    # Header with co-PI
    add_header(doc, config)

    # Approval and study period
    tbl_cover = doc.add_table(rows=3, cols=2)
    tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl_cover.rows[0].cells[0], "核准日期", bold=True, size=11)
    add_ct(tbl_cover.rows[0].cells[1], dates.get("irb_approval_date", ""), size=11)
    add_ct(tbl_cover.rows[1].cells[0], "研究期間", bold=True, size=11)
    add_ct(tbl_cover.rows[1].cells[1],
           f"{dates.get('study_start', '')} 至 {dates.get('study_end', '')}", size=11)
    add_ct(tbl_cover.rows[2].cells[0], "報告日期", bold=True, size=11)
    add_ct(tbl_cover.rows[2].cells[1], "＿＿＿＿年＿＿月＿＿日", size=11)
    apply_tb(tbl_cover)

    doc.add_page_break()

    # ---- Table of contents ----
    add_p(doc, "目　錄", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    sections = [
        ("一", "研究背景"),
        ("二", "研究目的"),
        ("三", "研究設計"),
        ("四", "研究參與者"),
        ("五", "研究方法"),
        ("六", "研究結果分析與討論"),
        ("七", "結論"),
        ("八", "參考文獻"),
    ]

    for num, title in sections:
        add_p(doc, f"{num}、{title}", size=12, sa=Pt(6))

    doc.add_page_break()

    # ---- Content sections ----
    for num, title in sections:
        add_p(doc, f"{num}、{title}", bold=True, size=14, sa=Pt(6), sb=Pt(12))
        add_p(doc, "（請填寫本節內容）", size=12, sa=Pt(12))
        doc.add_paragraph()

    # Footer
    add_footer(doc, "3", "IRB.SF038", "2022/09/01")

    path = os.path.join(output_dir, "SF038_結案報告書.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF023 — 資料及安全性監測計畫報告書
# ---------------------------------------------------------------------------


def generate_sf023(config, output_dir):
    """Generate SF023: Data Safety Monitoring Plan report."""
    doc = init_doc(sz=12)
    study = config["study"]
    subjects = config["subjects"]
    is_retro = study.get("type", "") == "retrospective"

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "資料及安全性監測計畫報告書", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header
    add_header(doc, config)

    # ---- 一、受試者人數 ----
    add_p(doc, "一、受試者人數", bold=True, size=12, sa=Pt(6))
    tbl1 = doc.add_table(rows=2, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "預計收案人數", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1], str(subjects.get("planned_n", 0)), size=10)
    add_ct(tbl1.rows[1].cells[0], "實際收案人數", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1], str(subjects.get("actual_n", 0)), size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、試驗藥品 ----
    add_p(doc, "二、試驗藥品", bold=True, size=12, sa=Pt(6))
    is_drug = study.get("drug_device", False)
    add_p(doc, f"{check(is_retro)}不適用（回溯性研究）", size=10, sa=Pt(2))
    add_p(doc, f"{check(is_drug)}有試驗藥品", size=10, sa=Pt(2))
    add_p(doc, f"{check(not is_retro and not is_drug)}無試驗藥品", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 三、計畫類別 ----
    add_p(doc, "三、計畫類別", bold=True, size=12, sa=Pt(6))
    type_map = {
        "retrospective": "回溯性研究",
        "prospective": "前瞻性研究",
        "clinical_trial": "臨床試驗",
        "genetic": "基因研究",
        "multicenter": "多中心研究",
    }
    study_type = study.get("type", "retrospective")
    add_p(doc, type_map.get(study_type, study_type), size=10, sa=Pt(4))

    design_map = {
        "cohort": "世代研究",
        "case_control": "病例對照研究",
        "cross_sectional": "橫斷面研究",
        "rct": "隨機對照試驗",
        "single_arm": "單臂研究",
    }
    study_design = study.get("design", "")
    if study_design:
        add_p(doc, f"研究設計：{design_map.get(study_design, study_design)}",
              size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 四、是否涉及易受傷害族群 ----
    add_p(doc, "四、是否涉及易受傷害族群", bold=True, size=12, sa=Pt(6))
    vuln = subjects.get("vulnerable_population", False)
    add_p(doc, f"{check(vuln)}是  {check(not vuln)}否", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 五、計畫執行成效 ----
    add_p(doc, "五、計畫執行成效", bold=True, size=12, sa=Pt(6))

    # Monitoring table: header + 3 phase rows, 4 columns
    monitor_cols = ["監測階段", "監測日期", "監測結果", "改善措施"]
    phases = ["計畫啟動期", "計畫執行期", "計畫結案期"]
    tbl5 = doc.add_table(rows=1 + len(phases), cols=4)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, col_name in enumerate(monitor_cols):
        cell = tbl5.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    for ri, phase_name in enumerate(phases, start=1):
        add_ct(tbl5.rows[ri].cells[0], phase_name, size=10)
        add_ct(tbl5.rows[ri].cells[1], "", size=10)
        add_ct(tbl5.rows[ri].cells[2], "", size=10)
        add_ct(tbl5.rows[ri].cells[3], "", size=10)

    apply_tb(tbl5)
    doc.add_paragraph()

    # ---- 六、期中分析 ----
    add_p(doc, "六、期中分析", bold=True, size=12, sa=Pt(6))
    add_p(doc, f"{check(is_retro)}不適用", size=10, sa=Pt(2))
    add_p(doc, f"{check(False)}有，分析結果如下：", size=10, sa=Pt(2))
    add_p(doc, f"{check(False)}無", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 七、DSMB ----
    add_p(doc, "七、DSMB（資料及安全性監測委員會）", bold=True, size=12, sa=Pt(6))
    add_p(doc, f"{check(is_retro)}不適用", size=10, sa=Pt(2))
    add_p(doc, f"{check(False)}有設立DSMB", size=10, sa=Pt(2))
    add_p(doc, f"{check(False)}無", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 八、其它重大事件 ----
    add_p(doc, "八、其它重大事件", bold=True, size=12, sa=Pt(6))
    add_p(doc, f"{check(False)}有，說明如下：", size=10, sa=Pt(2))
    add_p(doc, f"{check(True)}無", size=10, sa=Pt(4))
    doc.add_paragraph()

    # PI signature block
    add_p(doc, "計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "2", "IRB.SF023", "2016/11/07")

    path = os.path.join(output_dir, "SF023_資料及安全性監測計畫報告書.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF036": generate_sf036,
    "SF037": generate_sf037,
    "SF038": generate_sf038,
    "SF023": generate_sf023,
}
