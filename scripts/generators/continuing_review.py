"""Continuing/interim review (期中審查) form generators for KFSYSCC IRB.

Generates SF030, SF031, SF032 forms from config dict.
"""
import os

from scripts.docx_utils import *
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# SF030 — 期中審查送審資料表 v8, 2025.03.03
# ---------------------------------------------------------------------------


def generate_sf030(config, output_dir):
    """Generate SF030: continuing review submission checklist."""
    doc = init_doc(sz=12)
    is_drug = config["study"].get("drug_device", False)
    cr = config.get("continuing_review", {})
    ext = cr.get("extension_requested", False)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "期中審查送審資料表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Admin line
    add_p(doc, "承辦人員：＿＿＿＿＿＿＿＿　收件日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(6))

    # Checklist table: 5 columns
    cols = ["備妥", "※必備", "資料項目", "備註", "IRB檢核"]
    items = [
        (True,  True,  "送審文件電子檔", ""),
        (True,  True,  "送審資料表及文件繳交完成簽收表", ""),
        (True,  True,  "期中報告書（主持人簽名及日期）", "IRB.SF031"),
        (True,  True,  "中文計畫摘要（本會核准版本）", ""),
        (ext,   False, "計畫展延申請表", "IRB.SF032"),
        (is_drug, False, "個案報告表", ""),
        (is_drug, False, "藥品使用紀錄", ""),
        (True,  True,  "資料及安全性監測計畫報告書", "IRB.SF023"),
        (is_drug, False, "顯著財務利益申報表", "IRB.SF094"),
        (is_drug, False, "臨床試驗/研究許可證明", "IRB.SF011"),
    ]

    tbl = doc.add_table(rows=1 + len(items), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, col_name in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        add_ct(cell, col_name, bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for ri, (required, mandatory, name, note) in enumerate(items, start=1):
        add_ct(tbl.rows[ri].cells[0], check(required), size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[1], "※" if mandatory else "", size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[ri].cells[2], name, size=10)
        add_ct(tbl.rows[ri].cells[3], note, size=10)
        add_ct(tbl.rows[ri].cells[4], "", size=10)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "8", "IRB.SF030", "2025/03/03")

    path = os.path.join(output_dir, "SF030_期中審查送審資料表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF031 — 期中報告書 v7, 2025.03.03
# ---------------------------------------------------------------------------


def generate_sf031(config, output_dir):
    """Generate SF031: continuing review report."""
    doc = init_doc(sz=12)
    study = config["study"]
    pi = config["pi"]
    dates = config["dates"]
    subjects = config["subjects"]
    cr = config.get("continuing_review", {})
    is_retro = study.get("type", "") == "retrospective"
    is_drug = study.get("drug_device", False)
    waiver = subjects.get("consent_waiver", False)

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "期中報告書", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header
    add_header(doc, config)

    # ---- 一、計畫期間 ----
    add_p(doc, "一、計畫期間", bold=True, size=12, sa=Pt(6))

    ext_requested = cr.get("extension_requested", False)
    has_amendments = cr.get("has_amendments", False)

    tbl1 = doc.add_table(rows=3, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "委員會核准計畫期間", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1],
           f"{dates.get('study_start', '')} 至 {dates.get('study_end', '')}",
           size=10)
    add_ct(tbl1.rows[1].cells[0], "是否曾提出展延", bold=True, size=10)
    add_ct(tbl1.rows[1].cells[1],
           f"{check(ext_requested)}是  {check(not ext_requested)}否", size=10)
    add_ct(tbl1.rows[2].cells[0], "是否曾提出修正案", bold=True, size=10)
    add_ct(tbl1.rows[2].cells[1],
           f"{check(has_amendments)}是  {check(not has_amendments)}否", size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、收案進度 ----
    add_p(doc, "二、收案進度", bold=True, size=12, sa=Pt(6))

    planned_n = subjects.get("planned_n", 0)
    actual_n = subjects.get("actual_n", 0)
    enrollment_status = cr.get("enrollment_status", "")

    tbl2 = doc.add_table(rows=3, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl2.rows[0].cells[0], "預計收錄人數", bold=True, size=10)
    add_ct(tbl2.rows[0].cells[1], str(planned_n), size=10)
    add_ct(tbl2.rows[1].cells[0], "目前收錄人數", bold=True, size=10)
    add_ct(tbl2.rows[1].cells[1], str(actual_n), size=10)
    add_ct(tbl2.rows[2].cells[0], "收案狀態", bold=True, size=10)
    status_text = (
        f"{check(enrollment_status == 'enrolling')}收案中  "
        f"{check(enrollment_status == 'closed')}停止收案  "
        f"{check(enrollment_status == 'analyzing')}資料分析中"
    )
    add_ct(tbl2.rows[2].cells[1], status_text, size=10)
    apply_tb(tbl2)

    # Group breakdown
    groups = subjects.get("groups", [])
    if groups:
        doc.add_paragraph()
        tbl_g = doc.add_table(rows=1 + len(groups), cols=2)
        tbl_g.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_ct(tbl_g.rows[0].cells[0], "組別", bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl_g.rows[0].cells[1], "人數", bold=True, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(tbl_g.rows[0].cells[0], "D9E2F3")
        set_cell_shading(tbl_g.rows[0].cells[1], "D9E2F3")
        for gi, grp in enumerate(groups, start=1):
            add_ct(tbl_g.rows[gi].cells[0], grp.get("name", ""), size=10)
            add_ct(tbl_g.rows[gi].cells[1], str(grp.get("n", 0)), size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
        apply_tb(tbl_g)
    doc.add_paragraph()

    # ---- 三、研究執行概況 ----
    add_p(doc, "三、研究執行概況", bold=True, size=12, sa=Pt(6))
    add_p(doc, "（請簡述本研究之執行進度與重要發現）", size=10, sa=Pt(4))
    doc.add_paragraph()
    doc.add_paragraph()

    # ---- 四、不良事件摘要 ----
    add_p(doc, "四、不良事件摘要", bold=True, size=12, sa=Pt(6))
    if is_retro:
        add_p(doc, "■不適用（回溯性研究）", size=10, sa=Pt(4))
    else:
        ae_cols = ["不良事件類別", "件數", "與研究之相關性", "處理方式"]
        ae_categories = [
            "嚴重不良事件 (SAE)",
            "非預期嚴重不良反應 (SUSAR)",
            "一般不良事件 (AE)",
        ]
        tbl_ae = doc.add_table(rows=1 + len(ae_categories), cols=4)
        tbl_ae.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, col_name in enumerate(ae_cols):
            cell = tbl_ae.rows[0].cells[ci]
            add_ct(cell, col_name, bold=True, size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(cell, "D9E2F3")
        for ri, cat in enumerate(ae_categories, start=1):
            add_ct(tbl_ae.rows[ri].cells[0], cat, size=10)
            add_ct(tbl_ae.rows[ri].cells[1], "", size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_ct(tbl_ae.rows[ri].cells[2], "", size=10)
            add_ct(tbl_ae.rows[ri].cells[3], "", size=10)
        apply_tb(tbl_ae)
    doc.add_paragraph()

    # ---- 五、偏離計畫書事件 ----
    add_p(doc, "五、偏離計畫書事件", bold=True, size=12, sa=Pt(6))
    deviations = cr.get("deviations", 0)
    add_p(doc, f"偏離次數：{deviations}", size=10, sa=Pt(4))
    if deviations == 0:
        add_p(doc, "本研究期間無偏離計畫書事件", size=10, sa=Pt(4))
    else:
        add_p(doc, "（請詳述偏離事件內容及處理方式）", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 六、知情同意執行狀況 ----
    add_p(doc, "六、知情同意執行狀況", bold=True, size=12, sa=Pt(6))
    add_p(doc,
          f"{check(waiver)}免取得知情同意  "
          f"{check(not waiver)}已依規定取得受試者知情同意",
          size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 七、資料安全維護 ----
    add_p(doc, "七、資料安全維護", bold=True, size=12, sa=Pt(6))
    data_safety = cr.get("data_safety", config.get("closure", {}).get("data_safety", {}))
    tbl7 = doc.add_table(rows=4, cols=2)
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    items_7 = [
        ("資料去識別化",
         f"{check(data_safety.get('deidentified', False))}是  "
         f"{check(not data_safety.get('deidentified', False))}否"),
        ("資料加密保存",
         f"{check(data_safety.get('encrypted', False))}是  "
         f"{check(not data_safety.get('encrypted', False))}否"),
        ("資料保存年限", f"{data_safety.get('retention_years', 7)} 年"),
        ("資料存取授權人員", data_safety.get("authorized_personnel", "")),
    ]
    for ri, (label, val) in enumerate(items_7):
        add_ct(tbl7.rows[ri].cells[0], label, bold=True, size=10)
        add_ct(tbl7.rows[ri].cells[1], val, size=10)
    apply_tb(tbl7)
    doc.add_paragraph()

    # ---- 八、主持人聲明 ----
    add_p(doc, "八、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明以上所述均屬實，本計畫之執行均遵守研究計畫書及相關法規之規定。",
          size=10, sa=Pt(4))
    add_p(doc, f"計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "7", "IRB.SF031", "2025/03/03")

    path = os.path.join(output_dir, "SF031_期中報告書.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF032 — 計畫展延申請表 v2, 2016.11.07
# ---------------------------------------------------------------------------


def generate_sf032(config, output_dir):
    """Generate SF032: study extension request form."""
    doc = init_doc(sz=12)
    dates = config["dates"]
    subjects = config["subjects"]

    # Title
    add_p(doc, "和信治癌中心醫院 人體試驗委員會", bold=True, size=16,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(4))
    add_p(doc, "計畫展延申請表", bold=True, size=14,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header
    add_header(doc, config)

    # ---- 一、原核准期間 ----
    add_p(doc, "一、原核准期間", bold=True, size=12, sa=Pt(6))
    tbl1 = doc.add_table(rows=1, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl1.rows[0].cells[0], "委員會核准計畫期間", bold=True, size=10)
    add_ct(tbl1.rows[0].cells[1],
           f"{dates.get('study_start', '')} 至 {dates.get('study_end', '')}",
           size=10)
    apply_tb(tbl1)
    doc.add_paragraph()

    # ---- 二、申請展延期間 ----
    add_p(doc, "二、申請展延期間", bold=True, size=12, sa=Pt(6))
    add_p(doc, "自＿＿年＿＿月＿＿日　至＿＿年＿＿月＿＿日", size=10, sa=Pt(4))
    doc.add_paragraph()

    # ---- 三、展延原因 ----
    add_p(doc, "三、展延原因", bold=True, size=12, sa=Pt(6))
    add_p(doc, "（請詳述申請展延之原因）", size=10, sa=Pt(4))
    doc.add_paragraph()
    doc.add_paragraph()

    # ---- 四、目前收案進度 ----
    add_p(doc, "四、目前收案進度", bold=True, size=12, sa=Pt(6))
    planned_n = subjects.get("planned_n", 0)
    actual_n = subjects.get("actual_n", 0)
    tbl4 = doc.add_table(rows=2, cols=2)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl4.rows[0].cells[0], "預計收案人數", bold=True, size=10)
    add_ct(tbl4.rows[0].cells[1], str(planned_n), size=10)
    add_ct(tbl4.rows[1].cells[0], "實際收案人數", bold=True, size=10)
    add_ct(tbl4.rows[1].cells[1], str(actual_n), size=10)
    apply_tb(tbl4)
    doc.add_paragraph()

    # ---- 五、主持人聲明 ----
    add_p(doc, "五、主持人聲明", bold=True, size=12, sa=Pt(6))
    add_p(doc, "本人聲明以上所述均屬實，並承諾於展延期間內遵守研究計畫書及相關法規之規定。",
          size=10, sa=Pt(4))
    add_p(doc, f"計畫主持人簽名：＿＿＿＿＿＿＿＿＿＿　日期：＿＿＿＿年＿＿月＿＿日",
          size=10, sa=Pt(12))

    # Footer
    add_footer(doc, "2", "IRB.SF032", "2016/11/07")

    path = os.path.join(output_dir, "SF032_計畫展延申請表.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF030": generate_sf030,
    "SF031": generate_sf031,
    "SF032": generate_sf032,
}
