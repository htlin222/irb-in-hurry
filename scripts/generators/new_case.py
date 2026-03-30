"""New case submission form generators: SF001, SF002, SF094, SF011, SF022.

Generates DOCX forms for new IRB case submission at KFSYSCC.
"""

import os

from scripts.docx_utils import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm


# ---------------------------------------------------------------------------
# SF001 — 新案審查送審資料表 (v9, 2025.03.03)
# ---------------------------------------------------------------------------

def generate_sf001(config, output_dir):
    """Generate SF001: submission checklist for new case review."""
    doc = init_doc(sz=12)

    # Title block
    add_p(doc, "和信治癌中心醫院 人體試驗委員會",
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    add_p(doc, "新案審查送審資料表",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Administrative line
    add_p(doc, "承辦單位：人體試驗委員會　　　IRB承辦人：＿＿＿＿＿＿", size=11)

    # Checklist table
    study = config["study"]
    subjects = config.get("subjects", {})
    review_type = study.get("review_type", "")
    drug_device = study.get("drug_device", False)
    consent_waiver = subjects.get("consent_waiver", False)

    tbl = doc.add_table(rows=13, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["備妥", "藥品/非藥品", "資料項目", "備註", "IRB檢核"]
    for idx, h in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        add_ct(cell, h, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    # Checklist items: (備妥, 藥品/非藥品, 資料項目, 備註)
    items = [
        ("■", "※", "送審文件電子檔", "email至irb@kfsyscc.org"),
        ("■", "※", "送審資料表及文件繳交完成簽收表", ""),
        ("■", "※", "新案申請書（主持人簽名及日期）", "IRB.SF002"),
        ("■", "※", "中文計畫摘要", ""),
        ("■", "※", "研究計畫書", ""),
        ("■", "※", "主持人學經歷", "附件"),
        (check(drug_device), "", "受試者同意書",
         "IRB.SF062/SF063/SF090" if drug_device else ""),
        (check(not consent_waiver), "", "個案報告表（Case Report Form）", ""),
        ("■", "※", "顯著財務利益申報表", "IRB.SF094"),
        (check(review_type == "expedited"), "", "簡易審查範圍檢核表", "IRB.SF003"),
        (check(review_type == "exempt"), "", "免予審查範圍檢核表", "IRB.SF004"),
        (check(consent_waiver), "", "免取得知情同意檢核表", "IRB.SF005"),
    ]

    for row_idx, (ready, category, item, note) in enumerate(items, start=1):
        add_ct(tbl.rows[row_idx].cells[0], ready, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[row_idx].cells[1], category, size=10,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_ct(tbl.rows[row_idx].cells[2], item, size=10)
        add_ct(tbl.rows[row_idx].cells[3], note, size=9)
        add_ct(tbl.rows[row_idx].cells[4], "", size=10)  # IRB檢核 (blank)

    apply_tb(tbl)

    # Footer
    add_footer(doc, "9", "IRB.SF001", "2025/03/03")

    # Save
    fname = f"SF001_{config['study']['irb_no']}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF002 — 研究計畫申請書 (v11, 2025.03.03)
# ---------------------------------------------------------------------------

def generate_sf002(config, output_dir):
    """Generate SF002: main research application form."""
    doc = init_doc(sz=12)

    study = config["study"]
    pi = config["pi"]
    co_pis = config.get("co_pi", [])
    subjects = config.get("subjects", {})
    dates = config.get("dates", {})
    review_type = study.get("review_type", "")
    consent_waiver = subjects.get("consent_waiver", False)

    # Title block
    add_p(doc, "和信治癌中心醫院 人體試驗委員會",
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    add_p(doc, "研究計畫申請書",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # ── Section 一、計畫基本資料 ──
    add_p(doc, "一、計畫基本資料", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    tbl = doc.add_table(rows=6, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # IRB 編號
    add_ct(tbl.rows[0].cells[0], "KFSYSCC-IRB編號", bold=True, size=11)
    add_ct(tbl.rows[0].cells[1], study["irb_no"], size=11)

    # 計畫編號
    add_ct(tbl.rows[1].cells[0], "計畫編號", bold=True, size=11)
    add_ct(tbl.rows[1].cells[1], study.get("project_no", "不適用"), size=11)

    # 計畫名稱
    add_ct(tbl.rows[2].cells[0], "計畫名稱", bold=True, size=11)
    cell = tbl.rows[2].cells[1]
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f"（中文）{study['title_zh']}")
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(f"（英文）{study['title_en']}")
    set_run_font(run2, size=10)
    p2.paragraph_format.space_after = Pt(2)

    # 計畫類別
    study_type = study.get("type", "")
    type_text = (
        f"{check(study_type == 'observational')}觀察性研究　"
        f"{check(study_type == 'interventional')}介入性研究　"
        f"{check(study_type == 'clinical_trial')}臨床試驗　"
        f"{check(study_type == 'retrospective')}回溯性研究　"
        f"{check(study_type == 'other')}其他"
    )
    add_ct(tbl.rows[3].cells[0], "計畫類別", bold=True, size=11)
    add_ct(tbl.rows[3].cells[1], type_text, size=10)

    # 研究設計
    design = study.get("design", "")
    design_text = (
        f"{check(design == 'cohort')}世代研究　"
        f"{check(design == 'case_control')}病例對照　"
        f"{check(design == 'cross_sectional')}橫斷面　"
        f"{check(design == 'rct')}隨機對照　"
        f"{check(design == 'single_arm')}單臂試驗　"
        f"{check(design == 'case_series')}病例系列　"
        f"{check(design == 'other')}其他"
    )
    add_ct(tbl.rows[4].cells[0], "研究設計", bold=True, size=11)
    add_ct(tbl.rows[4].cells[1], design_text, size=10)

    # 多中心研究
    multicenter = study.get("multicenter", False)
    add_ct(tbl.rows[5].cells[0], "多中心研究", bold=True, size=11)
    add_ct(tbl.rows[5].cells[1],
           f"{check(multicenter)}是　{check(not multicenter)}否", size=10)

    apply_tb(tbl)
    doc.add_paragraph()

    # ── Section 二、計畫主持人資訊 ──
    add_p(doc, "二、計畫主持人資訊", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    pi_rows = 5
    tbl2 = doc.add_table(rows=pi_rows, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_ct(tbl2.rows[0].cells[0], "姓名", bold=True, size=11)
    add_ct(tbl2.rows[0].cells[1], pi["name"], size=11)

    add_ct(tbl2.rows[1].cells[0], "英文姓名", bold=True, size=11)
    add_ct(tbl2.rows[1].cells[1], pi.get("name_en", ""), size=11)

    add_ct(tbl2.rows[2].cells[0], "服務單位／職稱", bold=True, size=11)
    add_ct(tbl2.rows[2].cells[1], pi["dept"], size=11)

    add_ct(tbl2.rows[3].cells[0], "電話", bold=True, size=11)
    add_ct(tbl2.rows[3].cells[1], pi.get("phone", ""), size=11)

    add_ct(tbl2.rows[4].cells[0], "Email", bold=True, size=11)
    add_ct(tbl2.rows[4].cells[1], pi.get("email", ""), size=11)

    apply_tb(tbl2)

    # Co-PI section
    if co_pis:
        add_p(doc, "共同主持人：", bold=True, size=11, sa=Pt(6), sb=Pt(4))
        tbl_co = doc.add_table(rows=len(co_pis) + 1, cols=3)
        tbl_co.alignment = WD_TABLE_ALIGNMENT.CENTER

        co_headers = ["姓名", "英文姓名", "服務單位／職稱"]
        for idx, h in enumerate(co_headers):
            add_ct(tbl_co.rows[0].cells[idx], h, bold=True, size=10,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(tbl_co.rows[0].cells[idx], "D9E2F3")

        for ri, cp in enumerate(co_pis, start=1):
            add_ct(tbl_co.rows[ri].cells[0], cp.get("name", ""), size=10)
            add_ct(tbl_co.rows[ri].cells[1], cp.get("name_en", ""), size=10)
            add_ct(tbl_co.rows[ri].cells[2], cp.get("dept", ""), size=10)

        apply_tb(tbl_co)

    doc.add_paragraph()

    # ── Section 三、審查類別 ──
    add_p(doc, "三、審查類別", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    review_text = (
        f"{check(review_type == 'full_board')}一般審查　　"
        f"{check(review_type == 'expedited')}簡易審查　　"
        f"{check(review_type == 'exempt')}免予審查"
    )
    add_p(doc, review_text, size=11, sa=Pt(4))

    # ── Section 四、計畫期間 ──
    add_p(doc, "四、計畫期間", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    start = dates.get("study_start", "＿＿＿＿")
    end = dates.get("study_end", "＿＿＿＿")
    add_p(doc, f"預計起迄日期：{start} 至 {end}", size=11, sa=Pt(4))

    data_period = dates.get("data_period", "")
    if data_period:
        add_p(doc, f"資料期間：{data_period}", size=11, sa=Pt(4))

    # ── Section 五、受試者 ──
    add_p(doc, "五、受試者", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    planned_n = subjects.get("planned_n", "＿＿")
    add_p(doc, f"預計收錄人數：{planned_n}", size=11, sa=Pt(4))

    vuln = subjects.get("vulnerable_population", False)
    add_p(doc, f"是否涉及易受傷害族群：{check(vuln)}是　{check(not vuln)}否",
          size=11, sa=Pt(4))

    consent_text = (
        f"知情同意：{check(not consent_waiver)}需取得　"
        f"{check(consent_waiver)}免取得"
    )
    add_p(doc, consent_text, size=11, sa=Pt(4))

    # ── Section 六、經費來源 ──
    add_p(doc, "六、經費來源", bold=True, size=12, sa=Pt(6), sb=Pt(6))
    add_p(doc, "研究者自發性研究", size=11, sa=Pt(4))

    # ── Section 七、計畫主持人聲明 ──
    add_p(doc, "七、計畫主持人聲明", bold=True, size=12, sa=Pt(6), sb=Pt(6))

    declaration = (
        "本人聲明上述資料均屬實，本研究計畫將依據赫爾辛基宣言、"
        "人體研究法及相關法規執行，並遵守本院人體試驗委員會之規定。"
    )
    add_p(doc, declaration, size=11, sa=Pt(8))

    # Signature block
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl_sig.rows[0].cells[0], "計畫主持人簽名：", bold=True, size=11)
    add_ct(tbl_sig.rows[0].cells[1], "", size=11)
    add_ct(tbl_sig.rows[1].cells[0], "日期：", bold=True, size=11)
    add_ct(tbl_sig.rows[1].cells[1], "", size=11)
    apply_tb(tbl_sig)

    # Footer
    add_footer(doc, "11", "IRB.SF002", "2025/03/03")

    # Save
    fname = f"SF002_{config['study']['irb_no']}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF094 — 顯著財務利益申報表 (v1, 2023.05.01)
# ---------------------------------------------------------------------------

def generate_sf094(config, output_dir):
    """Generate SF094: financial interest disclosure form."""
    doc = init_doc(sz=12)

    # Title block
    add_p(doc, "和信治癌中心醫院 人體試驗委員會",
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    add_p(doc, "臨床研究人員顯著財務利益申報表",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    # Header block
    add_header(doc, config)

    # Declaration section
    add_p(doc, "本人聲明就本研究計畫，本人及配偶或受扶養親屬：",
          bold=True, size=11, sa=Pt(8), sb=Pt(6))

    # Default: 無 (most cases)
    has_financial = False  # default assumption

    add_p(doc, f"{check(not has_financial)} 無任何顯著財務利益需申報",
          size=11, sa=Pt(4))
    add_p(doc, f"{check(has_financial)} 有下列顯著財務利益需申報：",
          size=11, sa=Pt(4))

    # Sub-items (indented)
    sub_items = [
        "持股／股票選擇權",
        "顧問費／諮詢費",
        "智慧財產權",
        "其他",
    ]
    for item in sub_items:
        add_p(doc, f"　　□ {item}", size=11, sa=Pt(2))

    doc.add_paragraph()

    # Signature block
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_ct(tbl_sig.rows[0].cells[0], "申報人簽名：", bold=True, size=11)
    add_ct(tbl_sig.rows[0].cells[1], "", size=11)
    add_ct(tbl_sig.rows[1].cells[0], "日期：", bold=True, size=11)
    add_ct(tbl_sig.rows[1].cells[1], "", size=11)
    apply_tb(tbl_sig)

    # Footer
    add_footer(doc, "1", "IRB.SF094", "2023/05/01")

    # Save
    fname = f"SF094_{config['study']['irb_no']}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF011 — 臨床試驗／研究許可證明 (stub)
# ---------------------------------------------------------------------------

def generate_sf011(config, output_dir):
    """Generate SF011: placeholder clinical trial permit form."""
    doc = init_doc(sz=12)

    add_p(doc, "和信治癌中心醫院 人體試驗委員會",
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    add_p(doc, "臨床試驗／研究許可證明",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    add_header(doc, config)

    add_p(doc, "（本表內容待補充）", size=11, sa=Pt(12))

    fname = f"SF011_{config['study']['irb_no']}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# SF022 — 資料及安全性監測計畫 (stub)
# ---------------------------------------------------------------------------

def generate_sf022(config, output_dir):
    """Generate SF022: placeholder data safety monitoring plan form."""
    doc = init_doc(sz=12)

    add_p(doc, "和信治癌中心醫院 人體試驗委員會",
          bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(2))
    add_p(doc, "資料及安全性監測計畫",
          bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

    add_header(doc, config)

    add_p(doc, "（本表內容待補充）", size=11, sa=Pt(12))

    fname = f"SF022_{config['study']['irb_no']}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Generator registry
# ---------------------------------------------------------------------------

ALL_GENERATORS = {
    "SF001": generate_sf001,
    "SF002": generate_sf002,
    "SF094": generate_sf094,
    "SF011": generate_sf011,
    "SF022": generate_sf022,
}
